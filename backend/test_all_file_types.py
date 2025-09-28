#!/usr/bin/env python3
"""
Comprehensive test script for all supported file types in the FileProcessingService.
This script tests each file type to ensure they work correctly with the core concepts.
"""

import os
import sys
import asyncio
import tempfile
from pathlib import Path
import json
import csv
import xml.etree.ElementTree as ET
from datetime import datetime

# Add the backend directory to the Python path
backend_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, backend_dir)

from app.services.file_processing_service import FileProcessingService
from app.core.logger import logger

class FileTypeTester:
    """Test all supported file types in the FileProcessingService"""
    
    def __init__(self):
        self.test_dir = tempfile.mkdtemp(prefix="file_type_test_")
        self.results = {}
        self.db_session = None  # Mock database session
        
    def create_test_files(self):
        """Create test files for all supported formats"""
        logger.info(f"Creating test files in: {self.test_dir}")
        
        # Document formats
        self._create_txt_file()
        self._create_rtf_file()
        self._create_docx_file()
        self._create_pdf_file()
        
        # Spreadsheet formats
        self._create_csv_file()
        self._create_xlsx_file()
        self._create_xls_file()
        
        # Data formats
        self._create_json_file()
        self._create_xml_file()
        
        # Code formats
        self._create_python_file()
        self._create_javascript_file()
        self._create_html_file()
        self._create_css_file()
        
        # Image formats (create simple test images)
        self._create_test_image_files()
        
    def _create_txt_file(self):
        """Create a test TXT file with fillable content"""
        content = """RHETORICAL ANALYSIS WORKSHEET

Instructions: Complete this worksheet by filling in the blanks.

1. The force that attracts objects toward each other is called _____.
2. Albert _____ developed the theory of relativity.
3. Plants use _____ to create food from sunlight.
4. The largest ocean on Earth is the _____ Ocean.
5. A three-sided shape is called a _____.

Word Bank: gravity, Einstein, oxygen, photosynthesis, Pacific Ocean, triangle, Shakespeare, revolution, equator, polar bear, telescope, pharaoh, pyramids, hydrogen, continental drift, volcano, earthquake, glacier, quantum, atom bomb, virus, pathogen, planet, asteroid, internet, electricity, magnetism, economy, satellite, galaxy, telescope, revolution, democracy, republic, parliament, monarchy, polar bear, ice sheets, survival, pharaoh, egypt, ruler, pyramids, tombs, hydrogen, lightest, element, carbon, diamond, atom, quantum, subatomic, mechanics, continental drift, continents, move, volcano, lava, erupts, earthquake, shaking, surface, glacier, ice, moving, atomic bomb, bomb, wwii, destructive, virus, infectious, non-living, pathogen, disease, microscopic, bacteria, cell, protein, dna, evolution, internet, computers, worldwide, electricity, magnetism, economy, shakespeare, playwright, english, triangle, three, sides, asteroid, rocky, body
"""
        file_path = os.path.join(self.test_dir, "test_document.txt")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"Created TXT file: {file_path}")
        
    def _create_rtf_file(self):
        """Create a test RTF file"""
        content = """{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}} \\f0\\fs24 RHETORICAL ANALYSIS WORKSHEET\\par\\par Instructions: Complete this worksheet by filling in the blanks.\\par\\par 1. The force that attracts objects toward each other is called _____.\\par 2. Albert _____ developed the theory of relativity.\\par 3. Plants use _____ to create food from sunlight.\\par}"""
        file_path = os.path.join(self.test_dir, "test_document.rtf")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"Created RTF file: {file_path}")
        
    def _create_docx_file(self):
        """Create a test DOCX file using python-docx"""
        try:
            from docx import Document
            
            doc = Document()
            doc.add_heading('RHETORICAL ANALYSIS WORKSHEET', 0)
            doc.add_paragraph('Instructions: Complete this worksheet by filling in the blanks.')
            doc.add_paragraph('')
            
            # Add questions with blanks
            doc.add_paragraph('1. The force that attracts objects toward each other is called _____.')
            doc.add_paragraph('2. Albert _____ developed the theory of relativity.')
            doc.add_paragraph('3. Plants use _____ to create food from sunlight.')
            doc.add_paragraph('4. The largest ocean on Earth is the _____ Ocean.')
            doc.add_paragraph('5. A three-sided shape is called a _____.')
            
            # Add word bank
            doc.add_paragraph('')
            doc.add_paragraph('Word Bank: gravity, Einstein, oxygen, photosynthesis, Pacific Ocean, triangle, Shakespeare, revolution, equator, polar bear, telescope, pharaoh, pyramids, hydrogen, continental drift, volcano, earthquake, glacier, quantum, atom bomb, virus, pathogen, planet, asteroid, internet, electricity, magnetism, economy, satellite, galaxy, telescope, revolution, democracy, republic, parliament, monarchy, polar bear, ice sheets, survival, pharaoh, egypt, ruler, pyramids, tombs, hydrogen, lightest, element, carbon, diamond, atom, quantum, subatomic, mechanics, continental drift, continents, move, volcano, lava, erupts, earthquake, shaking, surface, glacier, ice, moving, atomic bomb, bomb, wwii, destructive, virus, infectious, non-living, pathogen, disease, microscopic, bacteria, cell, protein, dna, evolution, internet, computers, worldwide, electricity, magnetism, economy, shakespeare, playwright, english, triangle, three, sides, asteroid, rocky, body')
            
            file_path = os.path.join(self.test_dir, "test_document.docx")
            doc.save(file_path)
            logger.info(f"Created DOCX file: {file_path}")
            
        except ImportError:
            logger.warning("python-docx not available, skipping DOCX test")
            
    def _create_pdf_file(self):
        """Create a test PDF file"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            file_path = os.path.join(self.test_dir, "test_document.pdf")
            c = canvas.Canvas(file_path, pagesize=letter)
            width, height = letter
            
            # Add content
            c.drawString(100, height - 100, "RHETORICAL ANALYSIS WORKSHEET")
            c.drawString(100, height - 130, "Instructions: Complete this worksheet by filling in the blanks.")
            c.drawString(100, height - 160, "1. The force that attracts objects toward each other is called _____.")
            c.drawString(100, height - 180, "2. Albert _____ developed the theory of relativity.")
            c.drawString(100, height - 200, "3. Plants use _____ to create food from sunlight.")
            c.drawString(100, height - 220, "4. The largest ocean on Earth is the _____ Ocean.")
            c.drawString(100, height - 240, "5. A three-sided shape is called a _____.")
            
            c.save()
            logger.info(f"Created PDF file: {file_path}")
            
        except ImportError:
            logger.warning("reportlab not available, skipping PDF test")
            
    def _create_csv_file(self):
        """Create a test CSV file"""
        file_path = os.path.join(self.test_dir, "test_data.csv")
        with open(file_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Question', 'Answer', 'Category'])
            writer.writerow(['What force attracts objects?', '', 'Physics'])
            writer.writerow(['Who developed relativity?', '', 'Physics'])
            writer.writerow(['What do plants use for food?', '', 'Biology'])
            writer.writerow(['Largest ocean?', '', 'Geography'])
            writer.writerow(['Three-sided shape?', '', 'Math'])
        logger.info(f"Created CSV file: {file_path}")
        
    def _create_xlsx_file(self):
        """Create a test XLSX file"""
        try:
            from openpyxl import Workbook
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Test Sheet"
            
            # Add headers
            ws['A1'] = 'Question'
            ws['B1'] = 'Answer'
            ws['C1'] = 'Category'
            
            # Add data
            data = [
                ['What force attracts objects?', '', 'Physics'],
                ['Who developed relativity?', '', 'Physics'],
                ['What do plants use for food?', '', 'Biology'],
                ['Largest ocean?', '', 'Geography'],
                ['Three-sided shape?', '', 'Math']
            ]
            
            for row_idx, row_data in enumerate(data, 2):
                for col_idx, value in enumerate(row_data, 1):
                    ws.cell(row=row_idx, column=col_idx, value=value)
            
            file_path = os.path.join(self.test_dir, "test_data.xlsx")
            wb.save(file_path)
            logger.info(f"Created XLSX file: {file_path}")
            
        except ImportError:
            logger.warning("openpyxl not available, skipping XLSX test")
            
    def _create_xls_file(self):
        """Create a test XLS file"""
        try:
            import pandas as pd
            
            data = {
                'Question': ['What force attracts objects?', 'Who developed relativity?', 'What do plants use for food?', 'Largest ocean?', 'Three-sided shape?'],
                'Answer': ['', '', '', '', ''],
                'Category': ['Physics', 'Physics', 'Biology', 'Geography', 'Math']
            }
            
            df = pd.DataFrame(data)
            file_path = os.path.join(self.test_dir, "test_data.xls")
            df.to_excel(file_path, index=False)
            logger.info(f"Created XLS file: {file_path}")
            
        except ImportError:
            logger.warning("pandas not available, skipping XLS test")
            
    def _create_json_file(self):
        """Create a test JSON file"""
        data = {
            "assignment": {
                "title": "Science Quiz",
                "questions": [
                    {"question": "What force attracts objects?", "answer": "", "category": "Physics"},
                    {"question": "Who developed relativity?", "answer": "", "category": "Physics"},
                    {"question": "What do plants use for food?", "answer": "", "category": "Biology"},
                    {"question": "Largest ocean?", "answer": "", "category": "Geography"},
                    {"question": "Three-sided shape?", "answer": "", "category": "Math"}
                ],
                "word_bank": ["gravity", "Einstein", "photosynthesis", "Pacific Ocean", "triangle"]
            }
        }
        
        file_path = os.path.join(self.test_dir, "test_data.json")
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
        logger.info(f"Created JSON file: {file_path}")
        
    def _create_xml_file(self):
        """Create a test XML file"""
        root = ET.Element("assignment")
        root.set("title", "Science Quiz")
        
        questions = ET.SubElement(root, "questions")
        
        question_data = [
            {"question": "What force attracts objects?", "answer": "", "category": "Physics"},
            {"question": "Who developed relativity?", "answer": "", "category": "Physics"},
            {"question": "What do plants use for food?", "answer": "", "category": "Biology"},
            {"question": "Largest ocean?", "answer": "", "category": "Geography"},
            {"question": "Three-sided shape?", "answer": "", "category": "Math"}
        ]
        
        for q_data in question_data:
            question = ET.SubElement(questions, "question")
            question.set("category", q_data["category"])
            question.text = q_data["question"]
            
        word_bank = ET.SubElement(root, "word_bank")
        word_bank.text = "gravity, Einstein, photosynthesis, Pacific Ocean, triangle"
        
        file_path = os.path.join(self.test_dir, "test_data.xml")
        tree = ET.ElementTree(root)
        tree.write(file_path, encoding='utf-8', xml_declaration=True)
        logger.info(f"Created XML file: {file_path}")
        
    def _create_python_file(self):
        """Create a test Python file"""
        content = '''#!/usr/bin/env python3
"""
Test Python file for file processing
"""

def calculate_force(mass, acceleration):
    """
    Calculate force using F = ma
    TODO: Implement the calculation
    """
    pass

def main():
    """
    Main function
    TODO: Add main logic here
    """
    print("Hello World")
    # _____: Add more functionality here

if __name__ == "__main__":
    main()
'''
        file_path = os.path.join(self.test_dir, "test_code.py")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"Created Python file: {file_path}")
        
    def _create_javascript_file(self):
        """Create a test JavaScript file"""
        content = '''// Test JavaScript file for file processing

function calculateForce(mass, acceleration) {
    // TODO: Implement the calculation
    // F = ma
    return null;
}

function main() {
    console.log("Hello World");
    // _____: Add more functionality here
}

// Export functions
module.exports = {
    calculateForce,
    main
};
'''
        file_path = os.path.join(self.test_dir, "test_code.js")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"Created JavaScript file: {file_path}")
        
    def _create_html_file(self):
        """Create a test HTML file"""
        content = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Test HTML Document</title>
</head>
<body>
    <h1>Science Quiz</h1>
    <p>Complete the following questions:</p>
    
    <form>
        <label for="q1">What force attracts objects?</label>
        <input type="text" id="q1" name="q1" placeholder="_____">
        
        <label for="q2">Who developed relativity?</label>
        <input type="text" id="q2" name="q2" placeholder="_____">
        
        <label for="q3">What do plants use for food?</label>
        <input type="text" id="q3" name="q3" placeholder="_____">
        
        <!-- TODO: Add more questions -->
    </form>
</body>
</html>
'''
        file_path = os.path.join(self.test_dir, "test_code.html")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"Created HTML file: {file_path}")
        
    def _create_css_file(self):
        """Create a test CSS file"""
        content = '''/* Test CSS file for file processing */

body {
    font-family: Arial, sans-serif;
    /* TODO: Add more styling */
}

h1 {
    color: #333;
    /* _____: Add more heading styles */
}

.form {
    margin: 20px;
    /* TODO: Style the form */
}

/* TODO: Add responsive design */
'''
        file_path = os.path.join(self.test_dir, "test_code.css")
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"Created CSS file: {file_path}")
        
    def _create_test_image_files(self):
        """Create simple test image files"""
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            # Create a simple test image with text
            img = Image.new('RGB', (400, 200), color='white')
            draw = ImageDraw.Draw(img)
            
            # Add text to the image
            try:
                font = ImageFont.truetype("arial.ttf", 20)
            except:
                font = ImageFont.load_default()
            
            text = "Science Quiz\n1. What force attracts objects? _____\n2. Who developed relativity? _____\n3. What do plants use for food? _____"
            draw.text((10, 10), text, fill='black', font=font)
            
            # Save in different formats
            formats = ['png', 'jpg', 'gif', 'bmp', 'tiff', 'webp']
            for fmt in formats:
                file_path = os.path.join(self.test_dir, f"test_image.{fmt}")
                try:
                    if fmt == 'jpg':
                        img.save(file_path, 'JPEG')
                    elif fmt == 'tiff':
                        img.save(file_path, 'TIFF')
                    elif fmt == 'webp':
                        img.save(file_path, 'WEBP')
                    else:
                        img.save(file_path, fmt.upper())
                    logger.info(f"Created {fmt.upper()} file: {file_path}")
                except Exception as e:
                    logger.warning(f"Could not create {fmt} file: {e}")
                    
        except ImportError:
            logger.warning("PIL not available, skipping image tests")
            
    async def test_file_type(self, file_path: str, file_type: str):
        """Test a specific file type"""
        try:
            logger.info(f"Testing {file_type} file: {file_path}")
            
            # Create a mock database session
            class MockDBSession:
                pass
            
            service = FileProcessingService(MockDBSession())
            
            # Test analysis
            result = await service.process_file(file_path, user_id=1, action="analyze")
            
            self.results[file_type] = {
                'status': 'success',
                'file_path': file_path,
                'content_extracted': bool(result.get('content')),
                'fillable_sections_found': len(result.get('fillable_sections', [])),
                'analysis_completed': bool(result.get('analysis')),
                'file_type_detected': result.get('file_type'),
                'content_preview': str(result.get('content', {}).get('text', ''))[:200] + '...' if result.get('content', {}).get('text') else 'No text content',
                'fillable_sections': result.get('fillable_sections', [])[:3]  # First 3 sections
            }
            
            logger.info(f"✓ {file_type} test completed successfully")
            logger.info(f"  - Content extracted: {self.results[file_type]['content_extracted']}")
            logger.info(f"  - Fillable sections found: {self.results[file_type]['fillable_sections_found']}")
            
        except Exception as e:
            logger.error(f"✗ {file_type} test failed: {str(e)}")
            self.results[file_type] = {
                'status': 'failed',
                'file_path': file_path,
                'error': str(e)
            }
            
    async def run_all_tests(self):
        """Run tests for all file types"""
        logger.info("Starting comprehensive file type tests...")
        
        # Create test files
        self.create_test_files()
        
        # Define file types to test
        file_types = [
            ('txt', 'test_document.txt'),
            ('rtf', 'test_document.rtf'),
            ('docx', 'test_document.docx'),
            ('pdf', 'test_document.pdf'),
            ('csv', 'test_data.csv'),
            ('xlsx', 'test_data.xlsx'),
            ('xls', 'test_data.xls'),
            ('json', 'test_data.json'),
            ('xml', 'test_data.xml'),
            ('py', 'test_code.py'),
            ('js', 'test_code.js'),
            ('html', 'test_code.html'),
            ('css', 'test_code.css'),
            ('png', 'test_image.png'),
            ('jpg', 'test_image.jpg'),
            ('gif', 'test_image.gif'),
            ('bmp', 'test_image.bmp'),
            ('tiff', 'test_image.tiff'),
            ('webp', 'test_image.webp')
        ]
        
        # Test each file type
        for file_type, filename in file_types:
            file_path = os.path.join(self.test_dir, filename)
            if os.path.exists(file_path):
                await self.test_file_type(file_path, file_type)
            else:
                logger.warning(f"Test file not found: {file_path}")
                self.results[file_type] = {
                    'status': 'skipped',
                    'file_path': file_path,
                    'reason': 'File not created (missing dependencies)'
                }
        
        # Generate report
        self.generate_report()
        
    def generate_report(self):
        """Generate a comprehensive test report"""
        logger.info("\n" + "="*80)
        logger.info("FILE TYPE TESTING REPORT")
        logger.info("="*80)
        
        successful_tests = 0
        failed_tests = 0
        skipped_tests = 0
        
        for file_type, result in self.results.items():
            status = result['status']
            if status == 'success':
                successful_tests += 1
                logger.info(f"✓ {file_type.upper()}: SUCCESS")
                logger.info(f"  - Content extracted: {result['content_extracted']}")
                logger.info(f"  - Fillable sections: {result['fillable_sections_found']}")
                logger.info(f"  - Content preview: {result['content_preview']}")
                if result['fillable_sections']:
                    logger.info(f"  - Sample sections: {result['fillable_sections']}")
            elif status == 'failed':
                failed_tests += 1
                logger.info(f"✗ {file_type.upper()}: FAILED - {result['error']}")
            else:
                skipped_tests += 1
                logger.info(f"- {file_type.upper()}: SKIPPED - {result['reason']}")
        
        logger.info("\n" + "-"*80)
        logger.info(f"SUMMARY:")
        logger.info(f"  Successful: {successful_tests}")
        logger.info(f"  Failed: {failed_tests}")
        logger.info(f"  Skipped: {skipped_tests}")
        logger.info(f"  Total: {len(self.results)}")
        logger.info("-"*80)
        
        # Save detailed results to file
        report_file = os.path.join(self.test_dir, "test_report.json")
        with open(report_file, 'w', encoding='utf-8') as f:
            json.dump({
                'timestamp': datetime.utcnow().isoformat(),
                'summary': {
                    'successful': successful_tests,
                    'failed': failed_tests,
                    'skipped': skipped_tests,
                    'total': len(self.results)
                },
                'results': self.results
            }, f, indent=2)
        
        logger.info(f"Detailed report saved to: {report_file}")
        logger.info(f"Test files directory: {self.test_dir}")
        
    def cleanup(self):
        """Clean up test files"""
        import shutil
        try:
            shutil.rmtree(self.test_dir)
            logger.info(f"Cleaned up test directory: {self.test_dir}")
        except Exception as e:
            logger.warning(f"Could not clean up test directory: {e}")

async def main():
    """Main test function"""
    tester = FileTypeTester()
    
    try:
        await tester.run_all_tests()
    finally:
        # Keep test files for inspection
        logger.info(f"Test files preserved in: {tester.test_dir}")
        # tester.cleanup()

if __name__ == "__main__":
    asyncio.run(main())
