#!/usr/bin/env python3
"""
Manual Testing Script for File Completion System
This script helps you test the file completion functionality step by step.
"""
import os
import tempfile
import asyncio
from pathlib import Path
from docx import Document
from app.services.file_processing_service import FileProcessingService

class MockDB:
    """Mock database for testing"""
    def __init__(self):
        self.bind = None

def create_test_documents():
    """Create various test documents for testing"""
    test_docs = {}
    
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        
        # Test 1: Simple DOCX with placeholders
        print("📄 Creating Test Document 1: Simple DOCX with placeholders...")
        doc1 = Document()
        doc1.add_heading('RHETORICAL ANALYSIS TABLE', 0)
        doc1.add_paragraph('Instructions: Complete this rhetorical analysis table.')
        
        table1 = doc1.add_table(rows=4, cols=2)
        table1.cell(0, 0).text = 'Description'
        table1.cell(0, 1).text = 'Clues & Indicators'
        table1.cell(1, 0).text = 'WHY was it written? (purposes)'
        table1.cell(1, 1).text = '_____'
        table1.cell(2, 0).text = 'WHAT SPECIFIC CATEGORY of writing is this? (genre)'
        table1.cell(2, 1).text = '[INSERT ANSWER HERE]'
        table1.cell(3, 0).text = 'WHO was it written for? (audiences)?'
        table1.cell(3, 1).text = '_____'
        
        doc1_path = temp_path / "test_rhetorical_analysis.docx"
        doc1.save(doc1_path)
        test_docs['rhetorical_analysis'] = doc1_path
        
        # Test 2: Text document with placeholders
        print("📄 Creating Test Document 2: Text document with placeholders...")
        text_content = """Assignment: Complete this analysis

Section 1: Introduction
_____

Section 2: Main Points
[INSERT YOUR ANALYSIS HERE]

Section 3: Conclusion
_____

Instructions: Fill in the blank sections with your analysis."""
        
        text_path = temp_path / "test_analysis.txt"
        with open(text_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        test_docs['text_analysis'] = text_path
        
        # Test 3: Document with mixed content
        print("📄 Creating Test Document 3: Mixed content document...")
        doc3 = Document()
        doc3.add_heading('Project Report', 0)
        doc3.add_paragraph('This document contains various sections that need completion.')
        
        # Add a table
        table3 = doc3.add_table(rows=3, cols=3)
        table3.cell(0, 0).text = 'Task'
        table3.cell(0, 1).text = 'Status'
        table3.cell(0, 2).text = 'Notes'
        table3.cell(1, 0).text = 'Research'
        table3.cell(1, 1).text = '_____'
        table3.cell(1, 2).text = '[INSERT RESEARCH FINDINGS]'
        table3.cell(2, 0).text = 'Analysis'
        table3.cell(2, 1).text = 'In Progress'
        table3.cell(2, 2).text = '_____'
        
        doc3.add_paragraph('\nSummary:')
        doc3.add_paragraph('_____')
        
        doc3_path = temp_path / "test_project_report.docx"
        doc3.save(doc3_path)
        test_docs['project_report'] = doc3_path
        
        return test_docs, temp_path

async def test_file_completion(file_path, test_name, service):
    """Test file completion for a specific document"""
    print(f"\n🧪 Testing {test_name}...")
    print(f"📁 Input file: {file_path}")
    
    try:
        # Read the original file
        original_content = service._read_original_file(str(file_path))
        print(f"✅ Successfully read original file")
        
        # Create sample filled content
        filled_content = {
            'text': f"""This is AI-generated content for {test_name}. The analysis reveals important insights about the topic. Key findings include significant patterns and trends that warrant further investigation. The content demonstrates comprehensive understanding of the subject matter and provides valuable recommendations for future research and application."""
        }
        
        # Generate the filled file
        output_path = file_path.parent / f"completed_{file_path.name}"
        result_path = await service.generate_filled_file(
            original_path=str(file_path),
            filled_content=filled_content,
            output_path=str(output_path)
        )
        
        print(f"✅ Successfully generated completed file: {result_path}")
        
        # Verify the file
        if os.path.exists(result_path):
            file_size = os.path.getsize(result_path)
            print(f"📊 File size: {file_size} bytes")
            
            # Try to open it if it's a DOCX
            if result_path.suffix.lower() == '.docx':
                try:
                    test_doc = Document(result_path)
                    paragraph_count = len(test_doc.paragraphs)
                    table_count = len(test_doc.tables)
                    print(f"📄 Document has {paragraph_count} paragraphs and {table_count} tables")
                    
                    # Show preview
                    print("📋 Content preview:")
                    for i, paragraph in enumerate(test_doc.paragraphs[:3]):
                        if paragraph.text.strip():
                            preview = paragraph.text.strip()[:80]
                            print(f"  {i+1}. {preview}...")
                    
                except Exception as e:
                    print(f"❌ Error reading DOCX: {e}")
            
            return True
        else:
            print(f"❌ Output file does not exist: {result_path}")
            return False
            
    except Exception as e:
        print(f"❌ Test failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

async def run_manual_tests():
    """Run all manual tests"""
    print("🚀 Starting Manual File Completion Tests")
    print("=" * 50)
    
    # Create test documents
    test_docs, temp_dir = create_test_documents()
    
    # Initialize service
    service = FileProcessingService(MockDB())
    
    # Run tests
    results = {}
    
    for test_name, file_path in test_docs.items():
        success = await test_file_completion(file_path, test_name, service)
        results[test_name] = success
    
    # Summary
    print("\n" + "=" * 50)
    print("📊 TEST RESULTS SUMMARY")
    print("=" * 50)
    
    for test_name, success in results.items():
        status = "✅ PASSED" if success else "❌ FAILED"
        print(f"{test_name}: {status}")
    
    total_tests = len(results)
    passed_tests = sum(results.values())
    print(f"\nOverall: {passed_tests}/{total_tests} tests passed")
    
    if passed_tests == total_tests:
        print("🎉 All tests passed! The file completion system is working correctly.")
    else:
        print("⚠️  Some tests failed. Check the error messages above.")
    
    print(f"\n📁 Test files are in: {temp_dir}")
    print("💡 You can manually check the completed files to verify they look correct.")

def create_api_test_commands():
    """Generate API test commands for manual testing"""
    print("\n" + "=" * 50)
    print("🔗 API TESTING COMMANDS")
    print("=" * 50)
    
    print("1. Start the backend server:")
    print("   cd backend && py -m uvicorn app.main:app --reload")
    
    print("\n2. Test with curl (replace YOUR_TOKEN with actual token):")
    print("   curl -X POST 'http://localhost:8000/api/v1/file-processing/fill' \\")
    print("     -H 'Authorization: Bearer YOUR_TOKEN' \\")
    print("     -F 'file=@path/to/your/test_document.docx'")
    
    print("\n3. Or use the frontend:")
    print("   cd frontend && pnpm dev")
    print("   Open http://localhost:5173")

if __name__ == "__main__":
    print("Choose testing method:")
    print("1. Run automated tests")
    print("2. Show API testing commands")
    print("3. Both")
    
    choice = input("\nEnter choice (1-3): ").strip()
    
    if choice in ['1', '3']:
        asyncio.run(run_manual_tests())
    
    if choice in ['2', '3']:
        create_api_test_commands()
