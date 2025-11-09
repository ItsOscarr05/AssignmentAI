"""
File Write-Back Service for AssignmentAI
Implements answer insertion into original file structure while preserving layout and style per PRD requirements
"""
import os
import json
import tempfile
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
import logging

# Document processing libraries
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.shared import qn
from docx.enum.text import WD_COLOR_INDEX
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import black, blue

# Spreadsheet processing
import pandas as pd
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

# Image processing
from PIL import Image, ImageDraw, ImageFont

from app.core.logger import logger

class FileWriteBackService:
    """
    Service for writing answers back into original file structure
    Preserves layout, formatting, and document fidelity as per PRD requirements
    """
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self._write_back_pdf,
            'docx': self._write_back_docx,
            'doc': self._write_back_doc,
            'txt': self._write_back_txt,
            'rtf': self._write_back_rtf,
            'csv': self._write_back_csv,
            'xlsx': self._write_back_xlsx,
            'xls': self._write_back_xls,
            'py': self._write_back_code,
            'js': self._write_back_code,
            'java': self._write_back_code,
            'cpp': self._write_back_code,
            'c': self._write_back_code,
            'html': self._write_back_code,
            'css': self._write_back_code,
            'json': self._write_back_json,
            'xml': self._write_back_xml,
            'png': self._write_back_image,
            'jpg': self._write_back_image,
            'jpeg': self._write_back_image,
            'gif': self._write_back_image,
            'bmp': self._write_back_image,
            'tiff': self._write_back_image,
        }
    
    async def write_back_answers(
        self, 
        original_file_path: str, 
        filled_content: Dict[str, Any], 
        output_path: str,
        watermark: bool = False,
        subscription_tier: str = "free"
    ) -> str:
        """
        Write answers back into the original file structure
        Main entry point for file write-back operations
        """
        try:
            file_extension = Path(original_file_path).suffix[1:].lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format for write-back: {file_extension}")
            
            logger.info(f"Writing back answers to {file_extension} file: {output_path}")
            
            # Get the appropriate write-back handler
            write_back_handler = self.supported_formats[file_extension]
            
            # Perform write-back
            result_path = await write_back_handler(
                original_file_path, 
                filled_content, 
                output_path,
                watermark,
                subscription_tier
            )
            
            logger.info(f"Successfully wrote back answers to: {result_path}")
            return result_path
            
        except Exception as e:
            logger.error(f"Error writing back answers: {str(e)}")
            raise
    
    async def _write_back_docx(
        self, 
        original_path: str, 
        filled_content: Dict[str, Any], 
        output_path: str,
        watermark: bool = False,
        subscription_tier: str = "free"
    ) -> str:
        """Write back answers to DOCX files while preserving structure"""
        try:
            # Load the original document
            doc = Document(original_path)
            filled_text = filled_content.get('text', '')
            fillable_sections = filled_content.get('fillable_sections', [])
            
            # Process each fillable section
            section_index = 0
            filled_lines = filled_text.split('\n') if filled_text else []
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text
                
                # Check if this paragraph contains fillable content
                if self._is_fillable_paragraph(paragraph_text):
                    if section_index < len(filled_lines):
                        # Replace the paragraph content
                        self._replace_paragraph_content(paragraph, filled_lines[section_index])
                        section_index += 1
                
                # Add watermark to free tier users
                if watermark and subscription_tier == "free":
                    self._add_watermark_to_paragraph(paragraph)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if self._is_fillable_paragraph(paragraph.text):
                                if section_index < len(filled_lines):
                                    self._replace_paragraph_content(paragraph, filled_lines[section_index])
                                    section_index += 1
            
            # Add completion footer
            self._add_completion_footer(doc, subscription_tier, watermark)
            
            # Save the modified document
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            logger.error(f"Error writing back DOCX: {str(e)}")
            raise
    
    async def _write_back_pdf(
        self, 
        original_path: str, 
        filled_content: Dict[str, Any], 
        output_path: str,
        watermark: bool = False,
        subscription_tier: str = "free"
    ) -> str:
        """Write back answers to PDF files (append answer key)"""
        try:
            # For PDFs, we'll append an answer key as per PRD requirements
            filled_text = filled_content.get('text', '')
            
            # Create a new PDF with the answer key
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            # Add title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(100, height - 100, "AssignmentAI - Completed Assignment")
            c.drawString(100, height - 130, "=" * 50)
            
            # Add answers
            c.setFont("Helvetica", 12)
            y_position = height - 170
            
            lines = filled_text.split('\n')
            for line in lines:
                if y_position < 100:  # New page if needed
                    c.showPage()
                    y_position = height - 100
                
                # Truncate long lines
                if len(line) > 80:
                    line = line[:77] + "..."
                
                c.drawString(100, y_position, line)
                y_position -= 20
            
            # Add watermark if free tier
            if watermark and subscription_tier == "free":
                self._add_pdf_watermark(c, width, height)
            
            # Add footer
            c.setFont("Helvetica", 10)
            c.drawString(100, 50, f"Completed by AssignmentAI - {subscription_tier.title()} Plan")
            
            c.save()
            return output_path
            
        except Exception as e:
            logger.error(f"Error writing back PDF: {str(e)}")
            raise
    
    async def _write_back_txt(
        self, 
        original_path: str, 
        filled_content: Dict[str, Any], 
        output_path: str,
        watermark: bool = False,
        subscription_tier: str = "free"
    ) -> str:
        """Write back answers to TXT files with proper formatting"""
        try:
            # Read original content
            with open(original_path, 'r', encoding='utf-8') as f:
                original_text = f.read()
            
            # Build properly formatted document
            result_text = self._build_formatted_document(original_text, filled_content)
            
            # Add watermark if free tier
            if watermark and subscription_tier == "free":
                result_text += "\n\n" + "=" * 50
                result_text += "\nCompleted by AssignmentAI - Free Plan"
                result_text += "\nUpgrade to Plus/Pro for watermark-free output"
                result_text += "\n" + "=" * 50
            
            # Write result
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result_text)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error writing back TXT: {str(e)}")
            raise
    
    async def _write_back_rtf(
        self, 
        original_path: str, 
        filled_content: Dict[str, Any], 
        output_path: str,
        watermark: bool = False,
        subscription_tier: str = "free"
    ) -> str:
        """Write back answers to RTF files"""
        try:
            # Read original RTF content
            with open(original_path, 'r', encoding='utf-8') as f:
                original_rtf = f.read()
            
            filled_text = filled_content.get('text', '')
            filled_lines = filled_text.split('\n')
            
            # Replace fillable sections in RTF format
            result_rtf = self._replace_rtf_fillable_sections(original_rtf, filled_lines)
            
            # Add watermark if free tier
            if watermark and subscription_tier == "free":
                result_rtf += "\\par\\par " + "=" * 50
                result_rtf += "\\par Completed by AssignmentAI - Free Plan"
                result_rtf += "\\par Upgrade to Plus/Pro for watermark-free output"
                result_rtf += "\\par " + "=" * 50
            
            # Write result
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result_rtf)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error writing back RTF: {str(e)}")
            raise
    
    async def _write_back_csv(
        self, 
        original_path: str, 
        filled_content: Dict[str, Any], 
        output_path: str,
        watermark: bool = False,
        subscription_tier: str = "free"
    ) -> str:
        """Write back answers to CSV files"""
        try:
            # Load original CSV
            df = pd.read_csv(original_path)
            filled_data = filled_content.get('data', [])
            
            # Fill empty cells with filled data
            for idx, row_data in enumerate(filled_data):
                if idx < len(df):
                    for col in df.columns:
                        if col in row_data and (pd.isna(df.iloc[idx][col]) or df.iloc[idx][col] == ''):
                            df.iloc[idx][col] = row_data[col]
            
            # Add watermark row if free tier
            if watermark and subscription_tier == "free":
                watermark_row = pd.DataFrame([{
                    col: "Completed by AssignmentAI - Free Plan" if col == df.columns[0] else ""
                    for col in df.columns
                }])
                df = pd.concat([df, watermark_row], ignore_index=True)
            
            # Save result
            df.to_csv(output_path, index=False)
            return output_path
            
        except Exception as e:
            logger.error(f"Error writing back CSV: {str(e)}")
            raise
    
    async def _write_back_xlsx(
        self, 
        original_path: str, 
        filled_content: Dict[str, Any], 
        output_path: str,
        watermark: bool = False,
        subscription_tier: str = "free"
    ) -> str:
        """Write back answers to XLSX files"""
        try:
            # Load original workbook
            workbook = openpyxl.load_workbook(original_path)
            filled_sheets = filled_content.get('sheets', {})
            
            # Process each sheet
            for sheet_name in workbook.sheetnames:
                if sheet_name in filled_sheets:
                    worksheet = workbook[sheet_name]
                    sheet_data = filled_sheets[sheet_name]
                    
                    # Fill empty cells
                    for row_idx, row_data in enumerate(sheet_data, 1):
                        for col_idx, cell_value in enumerate(row_data, 1):
                            if cell_value and str(cell_value).strip():
                                cell = worksheet.cell(row=row_idx, column=col_idx)
                                if not cell.value or str(cell.value).strip() == '':
                                    cell.value = cell_value
                                    # Style filled cells
                                    cell.font = Font(color="0000FF")  # Blue text for filled cells
            
            # Add watermark sheet if free tier
            if watermark and subscription_tier == "free":
                watermark_sheet = workbook.create_sheet("AssignmentAI Info")
                watermark_sheet['A1'] = "Completed by AssignmentAI"
                watermark_sheet['A2'] = "Free Plan - Upgrade for watermark-free output"
                watermark_sheet['A1'].font = Font(bold=True, color="FF0000")
                watermark_sheet['A2'].font = Font(color="FF0000")
            
            # Save result
            workbook.save(output_path)
            return output_path
            
        except Exception as e:
            logger.error(f"Error writing back XLSX: {str(e)}")
            raise
    
    async def _write_back_xls(
        self, 
        original_path: str, 
        filled_content: Dict[str, Any], 
        output_path: str,
        watermark: bool = False,
        subscription_tier: str = "free"
    ) -> str:
        """Write back answers to XLS files"""
        try:
            # Load original XLS
            df_dict = pd.read_excel(original_path, sheet_name=None)
            filled_sheets = filled_content.get('sheets', {})
            
            # Process each sheet
            for sheet_name, df in df_dict.items():
                if sheet_name in filled_sheets:
                    filled_data = filled_sheets[sheet_name]
                    
                    # Fill empty cells
                    for idx, row_data in enumerate(filled_data):
                        if idx < len(df):
                            for col in df.columns:
                                if col in row_data and (pd.isna(df.iloc[idx][col]) or df.iloc[idx][col] == ''):
                                    df.iloc[idx][col] = row_data[col]
            
            # Add watermark sheet if free tier
            if watermark and subscription_tier == "free":
                watermark_df = pd.DataFrame([{
                    "AssignmentAI": "Completed by AssignmentAI - Free Plan",
                    "Upgrade": "Upgrade for watermark-free output"
                }])
                df_dict["AssignmentAI Info"] = watermark_df
            
            # Save result
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                for sheet_name, df in df_dict.items():
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error writing back XLS: {str(e)}")
            raise
    
    async def _write_back_code(
        self, 
        original_path: str, 
        filled_content: Dict[str, Any], 
        output_path: str,
        watermark: bool = False,
        subscription_tier: str = "free"
    ) -> str:
        """Write back answers to code files"""
        try:
            # Read original code
            with open(original_path, 'r', encoding='utf-8') as f:
                original_code = f.read()
            
            filled_text = filled_content.get('text', '')
            filled_lines = filled_text.split('\n')
            
            # Replace TODO comments and empty function bodies
            result_code = self._replace_code_fillable_sections(original_code, filled_lines)
            
            # Add watermark comment if free tier
            if watermark and subscription_tier == "free":
                result_code += "\n\n" + "# " + "=" * 50
                result_code += "\n# Completed by AssignmentAI - Free Plan"
                result_code += "\n# Upgrade to Plus/Pro for watermark-free output"
                result_code += "\n# " + "=" * 50
            
            # Write result
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result_code)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error writing back code: {str(e)}")
            raise
    
    async def _write_back_json(
        self, 
        original_path: str, 
        filled_content: Dict[str, Any], 
        output_path: str,
        watermark: bool = False,
        subscription_tier: str = "free"
    ) -> str:
        """Write back answers to JSON files"""
        try:
            # Load original JSON
            with open(original_path, 'r', encoding='utf-8') as f:
                original_data = json.load(f)
            
            filled_data = filled_content.get('data', {})
            
            # Merge filled data into original
            merged_data = self._merge_json_data(original_data, filled_data)
            
            # Add watermark if free tier
            if watermark and subscription_tier == "free":
                merged_data["_assignmentai"] = {
                    "completed_by": "AssignmentAI",
                    "plan": "Free",
                    "watermark": True,
                    "upgrade_message": "Upgrade to Plus/Pro for watermark-free output"
                }
            
            # Write result
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(merged_data, f, indent=2)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error writing back JSON: {str(e)}")
            raise
    
    async def _write_back_xml(
        self, 
        original_path: str, 
        filled_content: Dict[str, Any], 
        output_path: str,
        watermark: bool = False,
        subscription_tier: str = "free"
    ) -> str:
        """Write back answers to XML files"""
        try:
            # Read original XML
            with open(original_path, 'r', encoding='utf-8') as f:
                original_xml = f.read()
            
            filled_text = filled_content.get('text', '')
            filled_lines = filled_text.split('\n')
            
            # Replace fillable sections in XML
            result_xml = self._replace_xml_fillable_sections(original_xml, filled_lines)
            
            # Add watermark if free tier
            if watermark and subscription_tier == "free":
                watermark_element = f"""
    <!-- {50*'='} -->
    <!-- Completed by AssignmentAI - Free Plan -->
    <!-- Upgrade to Plus/Pro for watermark-free output -->
    <!-- {50*'='} -->"""
                result_xml = result_xml.replace('</', f'{watermark_element}\n</')
            
            # Write result
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result_xml)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error writing back XML: {str(e)}")
            raise
    
    async def _write_back_image(
        self, 
        original_path: str, 
        filled_content: Dict[str, Any], 
        output_path: str,
        watermark: bool = False,
        subscription_tier: str = "free"
    ) -> str:
        """Write back answers to image files (create text file with answers)"""
        try:
            # For images, create a text file with the answers
            output_path = output_path.replace(Path(output_path).suffix, '.txt')
            
            filled_text = filled_content.get('text', '')
            
            result_text = f"AssignmentAI - Completed Assignment\n"
            result_text += f"Original Image: {Path(original_path).name}\n"
            result_text += f"{'='*50}\n\n"
            result_text += filled_text
            
            # Add watermark if free tier
            if watermark and subscription_tier == "free":
                result_text += "\n\n" + "=" * 50
                result_text += "\nCompleted by AssignmentAI - Free Plan"
                result_text += "\nUpgrade to Plus/Pro for watermark-free output"
                result_text += "\n" + "=" * 50
            
            # Write result
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(result_text)
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error writing back image: {str(e)}")
            raise
    
    async def _write_back_doc(self, original_path: str, filled_content: Dict[str, Any], output_path: str, watermark: bool = False, subscription_tier: str = "free") -> str:
        """Write back answers to DOC files (convert to DOCX)"""
        # For legacy DOC files, we'll create a new DOCX file
        return await self._write_back_docx(original_path, filled_content, output_path.replace('.doc', '.docx'), watermark, subscription_tier)
    
    # Helper methods
    
    def _is_fillable_paragraph(self, text: str) -> bool:
        """Check if a paragraph contains fillable content"""
        return (
            '_____' in text or
            text.strip() == '' or
            '[INSERT' in text or
            'TODO' in text or
            text.strip() == '_____'
        )
    
    def _replace_paragraph_content(self, paragraph, new_content: str):
        """Replace paragraph content while preserving formatting"""
        # Clear existing content
        for run in paragraph.runs:
            run.clear()
        
        # Add new content
        if paragraph.runs:
            paragraph.runs[0].text = new_content
        else:
            paragraph.text = new_content
    
    def _add_watermark_to_paragraph(self, paragraph):
        """Add watermark to paragraph for free tier users"""
        if paragraph.runs:
            run = paragraph.runs[0]
            run.add_text(" [AssignmentAI Free]")
            run.font.color.rgb = RGBColor(200, 200, 200)  # Light gray
    
    def _add_completion_footer(self, doc: Document, subscription_tier: str, watermark: bool):
        """Add completion footer to document"""
        doc.add_paragraph("")
        footer = doc.add_paragraph(f"Completed by AssignmentAI - {subscription_tier.title()} Plan")
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        
        if watermark:
            watermark_para = doc.add_paragraph("Upgrade to Plus/Pro for watermark-free output")
            watermark_para.runs[0].font.size = Pt(8)
            watermark_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    
    def _add_pdf_watermark(self, canvas, width: float, height: float):
        """Add watermark to PDF"""
        canvas.saveState()
        canvas.setFont("Helvetica", 48)
        canvas.setFillColorRGB(0.9, 0.9, 0.9)
        canvas.rotate(45)
        canvas.drawString(width/2 - 100, height/2 - 100, "AssignmentAI Free")
        canvas.restoreState()
    
    def _replace_fillable_sections(self, original_text: str, filled_lines: List[str]) -> str:
        """Replace fillable sections in text"""
        lines = original_text.split('\n')
        result_lines = []
        filled_index = 0
        
        for line in lines:
            if self._is_fillable_paragraph(line):
                if filled_index < len(filled_lines):
                    result_lines.append(filled_lines[filled_index])
                    filled_index += 1
                else:
                    result_lines.append(line)
            else:
                result_lines.append(line)
        
        return '\n'.join(result_lines)
    
    def _replace_rtf_fillable_sections(self, original_rtf: str, filled_lines: List[str]) -> str:
        """Replace fillable sections in RTF format"""
        # Simple replacement for RTF - more sophisticated parsing could be added
        result = original_rtf
        filled_index = 0
        
        while '_____' in result and filled_index < len(filled_lines):
            result = result.replace('_____', filled_lines[filled_index], 1)
            filled_index += 1
        
        return result
    
    def _replace_code_fillable_sections(self, original_code: str, filled_lines: List[str]) -> str:
        """Replace fillable sections in code"""
        lines = original_code.split('\n')
        result_lines = []
        filled_index = 0
        
        for line in lines:
            if ('TODO' in line or 'FIXME' in line or 'pass' in line or 
                line.strip() == '' or '_____' in line):
                if filled_index < len(filled_lines):
                    result_lines.append(filled_lines[filled_index])
                    filled_index += 1
                else:
                    result_lines.append(line)
            else:
                result_lines.append(line)
        
        return '\n'.join(result_lines)
    
    def _replace_xml_fillable_sections(self, original_xml: str, filled_lines: List[str]) -> str:
        """Replace fillable sections in XML"""
        # Simple replacement for XML - more sophisticated parsing could be added
        result = original_xml
        filled_index = 0
        
        while '_____' in result and filled_index < len(filled_lines):
            result = result.replace('_____', filled_lines[filled_index], 1)
            filled_index += 1
        
        return result
    
    def _merge_json_data(self, original_data: Dict, filled_data: Dict) -> Dict:
        """Merge filled data into original JSON data"""
        merged = original_data.copy()
        
        for key, value in filled_data.items():
            if key not in merged or merged[key] == '' or merged[key] is None:
                merged[key] = value
        
        return merged
    
    def _build_formatted_document(self, original_text: str, filled_content: Dict[str, Any]) -> str:
        """Build a properly formatted document with structured answers"""
        try:
            # Parse the original document structure
            lines = original_text.split('\n')
            result_lines = []
            
            # Extract questions and answers from filled content
            filled_sections = filled_content.get('fillable_sections', [])
            answers_by_question = {}
            
            # Map answers to questions
            for section in filled_sections:
                question_text = section.get('text', '')
                filled_text = section.get('filled_text', '')
                
                # Find the question number/identifier
                question_id = self._extract_question_id(question_text)
                if question_id:
                    answers_by_question[question_id] = filled_text
            
            # Process each line of the original document
            current_question = None
            answer_added = False
            
            for line in lines:
                line = line.strip()
                
                # Check if this is a question line
                question_id = self._extract_question_id(line)
                if question_id:
                    current_question = question_id
                    answer_added = False
                    result_lines.append(line)  # Add the question
                    result_lines.append("")    # Add blank line
                    
                    # Add the answer if we have one
                    if question_id in answers_by_question:
                        answer = answers_by_question[question_id]
                        # Format the answer properly
                        formatted_answer = self._format_answer(answer, question_id)
                        result_lines.append(formatted_answer)
                        result_lines.append("")  # Add blank line after answer
                        answer_added = True
                    
                elif line and not line.startswith(('Course:', 'Assignment:', 'Q1)', 'Q2)', 'Q3)', 'Q4)', 'Q5)')):
                    # Regular content line (not a question)
                    result_lines.append(line)
                elif line.startswith(('Course:', 'Assignment:')):
                    # Header information
                    result_lines.append(line)
                    result_lines.append("")
                elif line == "":
                    # Empty line - only add if we haven't just added an answer
                    if not answer_added:
                        result_lines.append(line)
            
            # Join all lines with proper spacing
            result_text = '\n'.join(result_lines)
            
            # Clean up extra blank lines
            result_text = '\n'.join([line for line in result_text.split('\n') if line.strip() or line == ''])
            
            return result_text
            
        except Exception as e:
            logger.error(f"Error building formatted document: {str(e)}")
            # Fallback to simple replacement
            return self._replace_fillable_sections(original_text, filled_content.get('text', '').split('\n'))
    
    def _extract_question_id(self, text: str) -> Optional[str]:
        """Extract question identifier from text"""
        import re
        
        # Look for patterns like "Q1)", "Q2)", "1)", "2)", etc.
        patterns = [
            r'Q(\d+)\)',      # Q1), Q2), etc.
            r'(\d+)\)',       # 1), 2), etc.
            r'Question\s+(\d+)',  # Question 1, Question 2, etc.
            r'(\d+)\.\s',     # 1. , 2. , etc.
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return f"Q{match.group(1)}"
        
        return None
    
    def _format_answer(self, answer: str, question_id: str) -> str:
        """Format an answer with proper structure"""
        # Clean up the answer text
        answer = answer.strip()
        
        # Remove any existing "Answer:" prefix
        if answer.lower().startswith('answer:'):
            answer = answer[7:].strip()
        
        # Format based on question type
        if 'word count' in answer.lower() or '≈' in answer:
            # Word count question - format with word count
            return f"Answer: {answer}"
        elif len(answer.split()) > 20:
            # Long answer - format as paragraph
            return f"Answer: {answer}"
        else:
            # Short answer - format simply
            return f"Answer: {answer}"
