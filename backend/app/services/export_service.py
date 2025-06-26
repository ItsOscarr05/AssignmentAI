import io
from typing import Dict, Any
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.lib.colors import black, blue, gray
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import HTTPException
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

class ExportService:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Setup custom paragraph styles for better formatting."""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=black
        ))

        # Subtitle style
        self.styles.add(ParagraphStyle(
            name='CustomSubtitle',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceAfter=20,
            textColor=blue
        ))

        # Body text style
        self.styles.add(ParagraphStyle(
            name='CustomBody',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceAfter=12,
            alignment=TA_LEFT,
            textColor=black
        ))

        # Metadata style
        self.styles.add(ParagraphStyle(
            name='CustomMetadata',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            alignment=TA_LEFT,
            textColor=gray
        ))

    async def export_to_pdf(self, content: str, options: Dict[str, Any]) -> bytes:
        """
        Export content to PDF format.
        
        Args:
            content: The content to export
            options: Export options including title, metadata, etc.
            
        Returns:
            PDF file as bytes
        """
        try:
            # Create a buffer to store the PDF
            buffer = io.BytesIO()
            
            # Determine page size
            page_size = A4 if options.get('pageSize', 'a4').lower() == 'a4' else letter
            orientation = options.get('orientation', 'portrait')
            
            # Create the PDF document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=page_size,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )

            # Build the story (content)
            story = []
            
            # Add title
            title = options.get('customTitle', 'Assignment')
            story.append(Paragraph(title, self.styles['CustomTitle']))
            story.append(Spacer(1, 20))
            
            # Add metadata if requested
            if options.get('includeMetadata', True):
                metadata_text = f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
                if options.get('source'):
                    metadata_text += f"<br/>Source: {options['source']}"
                story.append(Paragraph(metadata_text, self.styles['CustomMetadata']))
                story.append(Spacer(1, 20))
            
            # Process content
            content_paragraphs = self._process_content_for_pdf(content)
            for paragraph in content_paragraphs:
                story.append(paragraph)
                story.append(Spacer(1, 12))
            
            # Build the PDF
            doc.build(story)
            
            # Get the PDF content
            buffer.seek(0)
            pdf_content = buffer.getvalue()
            buffer.close()
            
            return pdf_content
            
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")

    async def export_to_word(self, content: str, options: Dict[str, Any]) -> bytes:
        """
        Export content to Word document format.
        
        Args:
            content: The content to export
            options: Export options including title, metadata, etc.
            
        Returns:
            Word document as bytes
        """
        try:
            # Create a new Word document
            doc = Document()
            
            # Add title
            title = options.get('customTitle', 'Assignment')
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata if requested
            if options.get('includeMetadata', True):
                metadata_paragraph = doc.add_paragraph()
                metadata_paragraph.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
                if options.get('source'):
                    metadata_paragraph.add_run(f"\nSource: {options['source']}")
                metadata_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add some spacing
            doc.add_paragraph()
            
            # Process content
            content_paragraphs = self._process_content_for_word(content)
            for paragraph_text in content_paragraphs:
                if paragraph_text.strip():
                    p = doc.add_paragraph(paragraph_text)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            word_content = buffer.getvalue()
            buffer.close()
            
            return word_content
            
        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to generate Word document: {str(e)}")

    def _process_content_for_pdf(self, content: str) -> list:
        """
        Process content for PDF formatting.
        
        Args:
            content: Raw content string
            
        Returns:
            List of formatted paragraphs
        """
        paragraphs = []
        
        # Split content into paragraphs
        content_paragraphs = content.split('\n\n')
        
        for paragraph in content_paragraphs:
            paragraph = paragraph.strip()
            if paragraph:
                # Check if it's a heading (starts with # or is all caps)
                if paragraph.startswith('#') or paragraph.isupper():
                    # Remove # symbols and format as heading
                    clean_paragraph = paragraph.lstrip('#').strip()
                    paragraphs.append(Paragraph(clean_paragraph, self.styles['CustomSubtitle']))
                else:
                    # Regular paragraph
                    paragraphs.append(Paragraph(paragraph, self.styles['CustomBody']))
        
        return paragraphs

    def _process_content_for_word(self, content: str) -> list:
        """
        Process content for Word document formatting.
        
        Args:
            content: Raw content string
            
        Returns:
            List of paragraph texts
        """
        paragraphs = []
        
        # Split content into paragraphs
        content_paragraphs = content.split('\n\n')
        
        for paragraph in content_paragraphs:
            paragraph = paragraph.strip()
            if paragraph:
                # Remove markdown formatting for Word
                clean_paragraph = paragraph.lstrip('#').strip()
                paragraphs.append(clean_paragraph)
        
        return paragraphs

    async def export_to_google_docs(self, content: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """
        Prepare content for Google Docs export.
        This would typically integrate with Google Docs API.
        
        Args:
            content: The content to export
            options: Export options
            
        Returns:
            Dictionary with export information
        """
        try:
            # For now, we'll return the content formatted for Google Docs
            # In a real implementation, you'd use the Google Docs API
            
            formatted_content = self._format_for_google_docs(content, options)
            
            return {
                'content': formatted_content,
                'title': options.get('customTitle', 'Assignment'),
                'export_type': 'google-docs',
                'instructions': 'Copy the content below and paste it into Google Docs'
            }
            
        except Exception as e:
            logger.error(f"Error preparing Google Docs export: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to prepare Google Docs export: {str(e)}")

    def _format_for_google_docs(self, content: str, options: Dict[str, Any]) -> str:
        """
        Format content for Google Docs.
        
        Args:
            content: Raw content
            options: Export options
            
        Returns:
            Formatted content string
        """
        formatted_lines = []
        
        # Add title
        title = options.get('customTitle', 'Assignment')
        formatted_lines.append(f"# {title}")
        formatted_lines.append("")
        
        # Add metadata if requested
        if options.get('includeMetadata', True):
            formatted_lines.append(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            if options.get('source'):
                formatted_lines.append(f"Source: {options['source']}")
            formatted_lines.append("")
        
        # Add content
        formatted_lines.append(content)
        
        return '\n'.join(formatted_lines)

    async def get_export_formats(self) -> Dict[str, Any]:
        """
        Get available export formats and their capabilities.
        """
        return {
            'pdf': {
                'name': 'PDF Document',
                'description': 'Portable Document Format - Best for printing and sharing',
                'capabilities': ['customizable_layout', 'metadata_inclusion', 'page_settings'],
                'supported_options': ['pageSize', 'orientation', 'margins', 'includeMetadata', 'customTitle']
            },
            'docx': {
                'name': 'Word Document',
                'description': 'Microsoft Word Document - Editable format',
                'capabilities': ['editable', 'metadata_inclusion', 'formatting'],
                'supported_options': ['includeMetadata', 'customTitle']
            },
            'google-docs': {
                'name': 'Google Docs',
                'description': 'Cloud-based collaboration',
                'capabilities': ['collaboration', 'cloud_storage', 'real_time_editing'],
                'supported_options': ['includeMetadata', 'customTitle']
            }
        } 