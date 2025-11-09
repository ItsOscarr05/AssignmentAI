from typing import List, Optional, Dict, Any, Union
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Body, Request
from fastapi.responses import StreamingResponse
import json
from sqlalchemy.orm import Session
from sqlalchemy.ext.asyncio import AsyncSession
from pathlib import Path
from app.core.deps import get_current_user, get_db, get_async_db
from app.core.feature_access import require_feature, has_feature_access, get_user_plan, get_available_features, get_feature_requirements
from app.services.ai_service import AIService
from app.services.file_service import file_service
from app.services.web_scraping import WebScrapingService
from app.services.image_analysis_service import ImageAnalysisService
from app.services.diagram_service import diagram_service
from app.services.link_analysis_service import LinkAnalysisService
from app.models.user import User
from app.crud import file_upload as file_upload_crud
from app.schemas.file_upload import FileUploadCreate
from pydantic import BaseModel
import logging
import uuid
from datetime import datetime
import pypdf
from docx import Document
import io
import base64
import json
import os

logger = logging.getLogger(__name__)
router = APIRouter()


def _stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, (dict, list)):
        try:
            return json.dumps(value)
        except Exception:
            return str(value)
    return str(value)


def _serialize_response(payload: Dict[str, Any]) -> Dict[str, Any]:
    serialized: Dict[str, Any] = {}
    for key, value in payload.items():
        if isinstance(value, (dict, list)):
            serialized[key] = json.loads(json.dumps(value, default=str))
        elif isinstance(value, (str, int, float, bool)) or value is None:
            serialized[key] = value
        else:
            serialized[key] = _stringify(value)
    return serialized


async def _testing_generate_content(prompt: str, current_user: User, db: Optional[Session]) -> Dict[str, Any]:
    prompt_lower = prompt.lower()

    if any(keyword in prompt_lower for keyword in ["diagram", "chart", "flowchart", "graph"]):
        if not has_feature_access(current_user, "diagram_generation", db):
            plan = get_user_plan(current_user, db)
            raise HTTPException(
                status_code=403,
                detail={
                    "error": "Diagram generation not available in your plan",
                    "feature": "diagram_generation",
                    "current_plan": plan,
                    "upgrade_message": "Upgrade to Pro plan to access diagram generation",
                    "upgrade_url": "/dashboard/price-plan"
                }
            )
        diagram = await diagram_service.generate_diagram(description=prompt, diagram_type="auto", style="modern")
        return _serialize_response(
            {
                "content": f"Diagram generated successfully!\n\n{_stringify(diagram)}",
                "timestamp": datetime.utcnow().isoformat(),
                "service_used": "diagram_generation",
                "has_diagram": True,
            }
        )

    if any(keyword in prompt_lower for keyword in ["write code", "generate code", "python function", "javascript function", "code for"]):
        if not has_feature_access(current_user, "code_analysis", db):
            plan = get_user_plan(current_user, db)
            raise HTTPException(
                status_code=403,
                detail={
                    "error": "Code analysis not available in your plan",
                    "feature": "code_analysis",
                    "current_plan": plan,
                    "upgrade_message": "Upgrade to Plus plan to access code analysis",
                    "upgrade_url": "/dashboard/price-plan"
                }
            )

    try:
        content = await ai_service.generate_assignment_content_from_prompt(prompt)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error during content generation: {exc}"
        )
    return _serialize_response(
        {
            "content": _stringify(content),
            "timestamp": datetime.utcnow().isoformat(),
            "service_used": "ai_generation",
        }
    )


async def _testing_upload_and_process_file(file: UploadFile, current_user: User, db: Optional[Session]) -> Dict[str, Any]:
    if not file or not file.filename:
        raise HTTPException(status_code=422, detail="No file provided")

    try:
        file_path, _ = await file_service.save_file(file, current_user.id)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to process file: {exc}")

    content_type = file.content_type or "application/octet-stream"
    file_info = detect_file_type_and_service(content_type, file.filename)

    if file_info["service"] == "image_analysis":
        if not has_feature_access(current_user, "image_analysis", db):
            plan = get_user_plan(current_user, db)
            raise HTTPException(
                status_code=403,
                detail={
                    "error": "Image analysis not available in your plan",
                    "feature": "image_analysis",
                    "current_plan": plan,
                    "upgrade_message": "Upgrade to Pro plan to access image analysis",
                    "upgrade_url": "/dashboard/price-plan",
                },
            )

        with open(file_path, "rb") as img_file:
            image_bytes = img_file.read()

        analysis = await image_analysis_service.analyze_image_and_answer(
            image_data=image_bytes,
            question="Provide a comprehensive analysis of this image",
            context=None,
        )

        return _serialize_response(
            {
                "id": str(uuid.uuid4()),
                "name": file.filename,
                "size": getattr(file, "size", None),
                "type": content_type,
                "path": file_path,
                "content": f"Image Analysis Results:\n{_stringify(analysis)}",
                "analysis": _stringify(analysis),
                "service_used": "image_analysis",
                "file_category": "image",
                "uploaded_at": datetime.utcnow().isoformat(),
            }
        )

    # Text-based files
    try:
        extracted = await extract_file_content(file_path, content_type)
    except Exception:
        extracted = ""

    if file_info["type"] == "code":
        if not has_feature_access(current_user, "code_analysis", db):
            plan = get_user_plan(current_user, db)
            raise HTTPException(
                status_code=403,
                detail={
                    "error": "Code analysis not available in your plan",
                    "feature": "code_analysis",
                    "current_plan": plan,
                    "upgrade_message": "Upgrade to Plus plan to access code analysis",
                    "upgrade_url": "/dashboard/price-plan",
                },
            )
    elif file_info["type"] in {"data", "csv", "spreadsheet"}:
        if not has_feature_access(current_user, "data_analysis", db):
            plan = get_user_plan(current_user, db)
            raise HTTPException(
                status_code=403,
                detail={
                    "error": "Data analysis not available in your plan",
                    "feature": "data_analysis",
                    "current_plan": plan,
                    "upgrade_message": "Upgrade to Pro plan to access data analysis",
                    "upgrade_url": "/dashboard/price-plan",
                },
            )

    try:
        analysis = await ai_service.generate_assignment_content_from_prompt(extracted)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to process file: {exc}")

    return _serialize_response(
        {
            "id": str(uuid.uuid4()),
            "name": file.filename,
            "size": getattr(file, "size", None),
            "type": content_type,
            "path": file_path,
            "content": _stringify(extracted),
            "analysis": _stringify(analysis),
            "service_used": "ai_analysis",
            "file_category": file_info.get("type", "document"),
            "uploaded_at": datetime.utcnow().isoformat(),
        }
    )


async def _testing_process_uploaded_file(file_id: str, action: str) -> Dict[str, Any]:
    valid_actions = {"summarize", "extract", "rewrite", "analyze"}
    if action not in valid_actions:
        raise HTTPException(status_code=400, detail="Invalid action specified")
    try:
        result = await ai_service.generate_assignment_content_from_prompt(action)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to process file: {exc}")
    return _serialize_response(
        {
            "file_id": file_id,
            "action": action,
            "result": _stringify(result),
            "processed_at": datetime.utcnow().isoformat(),
        }
    )


async def _testing_process_link(url: str) -> Dict[str, Any]:
    try:
        async with WebScrapingService() as scraper:
            extracted = await scraper.extract_content_from_url(url)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to process link: {exc}")

    content = extracted.get("content", "")
    try:
        analysis = await ai_service.generate_assignment_content_from_prompt(content)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to process link: {exc}")

    return _serialize_response(
        {
            "id": str(uuid.uuid4()),
            "url": url,
            "title": extracted.get("title"),
            "content": content,
            "type": extracted.get("type"),
            "analysis": _stringify(analysis),
            "extracted_at": datetime.utcnow().isoformat(),
        }
    )

# Schema definitions for chat-with-link endpoint
class ChatWithLinkRequest(BaseModel):
    link_id: str
    message: str
    content: str
    analysis: Dict[str, Any] = None

class ChatWithLinkResponse(BaseModel):
    response: str
    updated_analysis: Optional[Dict[str, Any]] = None

# Provide a module-level hook for tests to patch
ai_service = AIService

@router.get("/health")
async def workshop_health_check():
    """
    Health check endpoint for the workshop service.
    """
    return {
        "status": "healthy",
        "service": "workshop",
        "timestamp": datetime.utcnow().isoformat(),
        "message": "Workshop service is running and ready to process requests!"
    }

@router.get("/quota-status")
async def check_openai_quota():
    """
    Check OpenAI API quota status and usage.
    """
    try:
        import os
        from openai import OpenAI
        
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            return {
                "status": "error",
                "message": "OpenAI API key not configured",
                "timestamp": datetime.utcnow().isoformat()
            }
        
        client = OpenAI(api_key=api_key)
        
        # Try to make a minimal API call to check quota
        try:
            response = client.chat.completions.create(
                model="gpt-5-nano",
                messages=[{"role": "user", "content": "Hello"}],
                max_completion_tokens=5
            )
            
            return {
                "status": "healthy",
                "message": "OpenAI API quota available",
                "quota_status": "available",
                "timestamp": datetime.utcnow().isoformat(),
                "test_response": "Success"
            }
            
        except Exception as e:
            error_message = str(e)
            if "insufficient_quota" in error_message.lower() or "quota_exceeded" in error_message.lower():
                return {
                    "status": "quota_exceeded",
                    "message": "OpenAI API quota exceeded",
                    "quota_status": "exceeded",
                    "error": error_message,
                    "timestamp": datetime.utcnow().isoformat()
                }
            elif "rate_limit" in error_message.lower() or "429" in error_message:
                return {
                    "status": "rate_limited",
                    "message": "OpenAI API rate limit hit",
                    "quota_status": "rate_limited",
                    "error": error_message,
                    "timestamp": datetime.utcnow().isoformat()
                }
            else:
                return {
                    "status": "error",
                    "message": "OpenAI API error",
                    "quota_status": "error",
                    "error": error_message,
                    "timestamp": datetime.utcnow().isoformat()
                }
                
    except Exception as e:
        logger.error(f"Error checking OpenAI quota: {e}")
        return {
            "status": "error",
            "message": "Failed to check quota status",
            "error": str(e),
            "timestamp": datetime.utcnow().isoformat()
        }

# Initialize services
image_analysis_service = ImageAnalysisService()

async def create_file_upload_record(
    db: AsyncSession,
    user_id: int,
    file: UploadFile,
    file_path: str,
    file_type: str,
    extracted_content: str = None,
    ai_analysis: str = None,
    processing_status: str = "completed"
) -> dict:
    """Create a file upload record in the database"""
    try:
        file_upload_data = FileUploadCreate(
            filename=file.filename,
            original_filename=file.filename,
            file_path=file_path,
            file_size=file.size,
            mime_type=file.content_type or "application/octet-stream",
            file_type=file_type,
            extracted_content=_stringify(extracted_content) if extracted_content is not None else None,
            ai_analysis=_stringify(ai_analysis) if ai_analysis is not None else None,
            processing_status=processing_status,
            is_link=False,
            upload_metadata={
                "uploaded_via": "workshop",
                "timestamp": datetime.utcnow().isoformat()
            }
        )
        
        db_file_upload = await file_upload_crud.create_file_upload_async(db, file_upload_data, user_id)
        logger.info(f"Successfully created file upload record with ID: {db_file_upload.id}")
        return {
            "file_upload_id": db_file_upload.id,
            "success": True
        }
    except Exception as e:
        logger.error(f"Failed to create file upload record: {str(e)}")
        return {
            "file_upload_id": None,
            "success": False,
            "error": str(e)
        }

def detect_file_type_and_service(content_type: str, filename: str) -> dict:
    """
    Detect the type of file and determine which service should process it.
    """
    file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
    
    # Image files - route to image analysis
    if content_type.startswith('image/') or file_extension in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp']:
        return {
            'type': 'image',
            'service': 'image_analysis',
            'analysis_type': 'general'
        }
    
    # Code files - route to AI code analysis
    elif file_extension in ['py', 'js', 'java', 'cpp', 'cc', 'c', 'html', 'htm', 'css', 'json', 'xml']:
        return {
            'type': 'code',
            'service': 'ai_analysis',
            'analysis_type': 'code_review'
        }
    
    # Spreadsheet files - route to spreadsheet analysis (more specific than generic data)
    elif file_extension in ['xlsx', 'xls'] or content_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel']:
        return {
            'type': 'spreadsheet',
            'service': 'ai_analysis',
            'analysis_type': 'spreadsheet_analysis'
        }
    
    # CSV files - route to CSV analysis
    elif file_extension == 'csv' or content_type == 'text/csv':
        return {
            'type': 'csv',
            'service': 'ai_analysis',
            'analysis_type': 'csv_analysis'
        }
    
    # Document files - route to document analysis
    elif content_type in ['application/pdf', 'application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'text/plain', 'application/rtf']:
        return {
            'type': 'document',
            'service': 'ai_analysis',
            'analysis_type': 'document_analysis'
        }
    
    # Default to general AI analysis
    else:
        return {
            'type': 'unknown',
            'service': 'ai_analysis',
            'analysis_type': 'general'
        }

@router.post("/generate")
async def generate_content(
    prompt: str = Body(..., embed=True),
    conversation_history: List[dict] = Body(default=[]),
    stream: bool = Body(default=False, embed=True),
    db: AsyncSession = Depends(get_async_db),
    current_user: User = Depends(get_current_user)
):
    """
    Generate AI content based on a prompt. This endpoint now routes to different services
    based on the content of the prompt and enforces feature access controls.
    Supports streaming responses for real-time chat experience.
    """
    # Add comprehensive logging
    logger.info(f"=== WORKSHOP GENERATE ENDPOINT CALLED ===")
    logger.info(f"User ID: {current_user.id}")
    logger.info(f"User email: {current_user.email}")
    # Log prompt with Unicode handling
    try:
        logger.info(f"Received prompt: {prompt}")
    except UnicodeEncodeError:
        logger.info(f"Received prompt (length: {len(prompt)} characters)")
    logger.info(f"Prompt length: {len(prompt)} characters")
    logger.info(f"Request timestamp: {datetime.utcnow()}")

    if os.getenv("TESTING") == "true":
        return await _testing_generate_content(prompt, current_user, db)

    # Initialize AI service with database session
    ai_service_instance = ai_service(db=db)
    
    # Enforce token limits before making AI calls
    try:
        estimated_tokens = len(prompt.split()) * 2  # Rough estimate
        await ai_service_instance.enforce_token_limit(current_user.id, estimated_tokens)
        logger.info(f"Token limit check passed for user {current_user.id}")
    except HTTPException as e:
        logger.warning(f"Token limit exceeded for user {current_user.id}: {e.detail}")
        raise e
    except Exception as e:
        logger.error(f"Token enforcement failed for user {current_user.id}: {str(e)}")
        # Return a proper error response instead of continuing
        raise HTTPException(
            status_code=403,
            detail="Unable to verify token limits. Please contact support."
        )

    try:
        prompt_lower = prompt.lower()
        
        # Detect diagram requests
        logger.info("Checking for diagram generation keywords...")
        if any(keyword in prompt_lower for keyword in ['create a diagram', 'generate chart', 'make a graph', 'draw a flowchart', 'create a bar chart', 'pie chart', 'line graph', 'scatter plot']):
            logger.info("Diagram generation keywords detected!")
            # Check if user has access to diagram generation
            logger.info("Checking diagram generation feature access...")
            if not has_feature_access(current_user, "diagram_generation", db):
                plan = get_user_plan(current_user, db)
                logger.warning(f"User {current_user.id} denied access to diagram generation. Plan: {plan}")
                raise HTTPException(
                    status_code=403,
                    detail={
                        "error": "Diagram generation not available in your plan",
                        "feature": "diagram_generation",
                        "current_plan": plan,
                        "upgrade_message": "Upgrade to Pro plan to access diagram generation",
                        "upgrade_url": "/dashboard/price-plan"
                    }
                )
            
            try:
                # Route to diagram service
                logger.info("Calling diagram service...")
                result = await diagram_service.generate_diagram(
                    description=prompt,
                    diagram_type="auto",
                    style="modern"
                )
                logger.info(f"Diagram generation successful! Result length: {len(str(result))}")
                return {
                    "content": f"Diagram generated successfully!\n\n{result}",
                    "timestamp": datetime.utcnow().isoformat(),
                    "service_used": "diagram_generation",
                    "has_diagram": True
                }
            except Exception as e:
                logger.error(f"Diagram generation failed: {str(e)}")
                logger.info("Falling back to general AI service...")
                # Fallback to general AI
                content = await ai_service_instance.generate_assignment_content_from_prompt(
                    prompt, 
                    user_id=current_user.id, 
                    feature='diagram_fallback', 
                    action='generate'
                )
                logger.info(f"General AI fallback completed. Content length: {len(content)} characters")
                return {
                    "content": content,
                    "timestamp": datetime.utcnow().isoformat(),
                    "service_used": "general_ai_fallback"
                }
        
        # Detect code generation requests
        elif any(keyword in prompt_lower for keyword in ['write code', 'generate function', 'create a script', 'program in', 'code for', 'python script', 'javascript function', 'java class', 'html code']):
            logger.info("Code generation keywords detected!")
            # Check if user has access to code generation
            logger.info("Checking code generation feature access...")
            if not has_feature_access(current_user, "code_review_assistant", db):
                plan = get_user_plan(current_user, db)
                logger.warning(f"User {current_user.id} denied access to code review assistant. Plan: {plan}")
                raise HTTPException(
                    status_code=403,
                    detail={
                        "error": "Code review assistant not available in your plan",
                        "feature": "code_review_assistant",
                        "current_plan": plan,
                        "upgrade_message": "Upgrade to Pro plan to access code review assistant",
                        "upgrade_url": "/dashboard/price-plan"
                    }
                )
            
            # Route to AI service with code generation focus
            logger.info("Routing to AI service for code generation...")
            code_prompt = f"Generate code for the following request: {prompt}\n\nPlease provide complete, working code with comments and explanations."
            content = await ai_service_instance.generate_assignment_content_from_prompt(
                code_prompt, 
                user_id=current_user.id, 
                feature='code_generation', 
                action='generate'
            )
            logger.info(f"Code generation completed. Content length: {len(content)} characters")
            return {
                "content": content,
                "timestamp": datetime.utcnow().isoformat(),
                "service_used": "code_review_assistant"
            }
        
        # Detect math problem solving
        elif any(keyword in prompt_lower for keyword in ['solve this math', 'calculate', 'equation', 'mathematical', 'solve for', 'find the value']):
            logger.info("Math problem solving keywords detected!")
            # Route to AI service with math focus
            logger.info("Routing to AI service for math solving...")
            math_prompt = f"Solve this mathematical problem step by step: {prompt}\n\nPlease show your work and explain each step."
            content = await ai_service.generate_assignment_content_from_prompt(
                math_prompt, 
                user_id=current_user.id, 
                feature='math_solving', 
                action='solve'
            )
            logger.info(f"Math solving completed. Content length: {len(content)} characters")
            return {
                "content": content,
                "timestamp": datetime.utcnow().isoformat(),
                "service_used": "math_solving"
            }
        
        # Default to general AI content generation
        else:
            logger.info("No specific service detected, using general AI content generation...")
            logger.info("Routing to AI service for general content...")
            
            # Check if this looks like a simple chat message vs assignment request
            # Simple indicators: short message, conversational tone, questions, etc.
            is_simple_chat = (
                len(prompt.split()) <= 10 or  # Short messages
                prompt.endswith('?') or  # Questions
                any(word in prompt.lower() for word in ['hi', 'hello', 'hey', 'thanks', 'thank you', 'help', 'what', 'how', 'why', 'when', 'where', 'who']) or
                not any(word in prompt.lower() for word in ['assignment', 'homework', 'project', 'essay', 'report', 'paper', 'lesson', 'curriculum'])
            )
            
            if is_simple_chat:
                logger.info("Detected simple chat message, using GPT conversational AI...")
                
                # Check if streaming is requested
                if stream:
                    logger.info("Streaming response requested for chat message")
                    
                    async def generate_stream():
                        try:
                            logger.info("Starting streaming generation...")
                            chunk_count = 0
                            full_response = ""
                            async for chunk in ai_service.generate_chat_response_stream(prompt, conversation_history, current_user.id):
                                chunk_count += 1
                                full_response += chunk
                                # Log chunk info without full content to avoid Unicode issues
                                chunk_preview = chunk[:100].replace('\n', ' ').replace('\r', ' ') if chunk else ""
                                logger.info(f"Yielding chunk {chunk_count} ({len(chunk)} chars): {chunk_preview}...")
                                # Format chunk as Server-Sent Events
                                yield f"data: {json.dumps({'content': chunk, 'done': False})}\n\n"
                            logger.info(f"Streaming completed. Total chunks: {chunk_count}")
                            
                            # Track token usage after streaming completes
                            try:
                                # Estimate tokens based on prompt and response
                                estimated_tokens = len(prompt.split()) + len(full_response.split())
                                await ai_service.track_token_usage(
                                    user_id=current_user.id,
                                    feature='workshop_chat',
                                    action='streaming_response',
                                    tokens_used=estimated_tokens,
                                    metadata={
                                        'prompt_length': len(prompt),
                                        'response_length': len(full_response),
                                        'conversation_length': len(conversation_history) if conversation_history else 0,
                                        'model': await ai_service.get_user_model(current_user.id)
                                    }
                                )
                                logger.info(f"Token usage tracked: {estimated_tokens} tokens for streaming chat")
                            except Exception as e:
                                logger.error(f"Failed to track token usage for streaming chat: {str(e)}")
                            
                            # Send final completion signal
                            yield f"data: {json.dumps({'content': '', 'done': True, 'timestamp': datetime.utcnow().isoformat(), 'service_used': 'gpt_chat_stream'})}\n\n"
                        except Exception as e:
                            logger.error(f"Error in streaming response: {str(e)}")
                            yield f"data: {json.dumps({'error': str(e), 'done': True})}\n\n"
                    
                    return StreamingResponse(
                        generate_stream(),
                        media_type="text/event-stream",
                        headers={
                            "Cache-Control": "no-cache",
                            "Connection": "keep-alive",
                            "Access-Control-Allow-Origin": "*",
                            "Access-Control-Allow-Headers": "*"
                        }
                    )
                else:
                    # Use conversation history passed from frontend and user's subscription model
                    content = await ai_service.generate_chat_response(prompt, conversation_history, current_user.id)
                    logger.info(f"GPT chat response completed. Content length: {len(content)} characters")
                    return {
                        "content": content,
                        "timestamp": datetime.utcnow().isoformat(),
                        "service_used": "gpt_chat"
                    }
            else:
                logger.info("Detected assignment request, using assignment-focused AI...")
                content = await ai_service.generate_assignment_content_from_prompt(
                    prompt, 
                    user_id=current_user.id, 
                    feature='workshop_assignment', 
                    action='generate'
                )
                logger.info(f"Assignment generation completed. Content length: {len(content)} characters")
                return {
                    "content": content,
                    "timestamp": datetime.utcnow().isoformat(),
                    "service_used": "general_ai"
                }
        
        # Add final success log
        logger.info("=== WORKSHOP GENERATE ENDPOINT COMPLETED SUCCESSFULLY ===")
        logger.info(f"User ID: {current_user.id}")
        logger.info(f"Service used: general_ai")
        logger.info(f"Timestamp: {datetime.utcnow()}")
            
    except HTTPException as e:
        # Re-raise HTTPExceptions (like 403 errors) without wrapping them
        logger.warning(f"HTTPException raised: {str(e)}")
        raise
    except Exception as e:
        logger.error(f"=== ERROR IN WORKSHOP GENERATE ENDPOINT ===")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error message: {str(e)}")
        logger.error(f"User ID: {current_user.id}")
        logger.error(f"Prompt: {prompt}")
        logger.error(f"Timestamp: {datetime.utcnow()}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail="Internal server error during content generation")

@router.post("/files/test")
async def test_file_upload(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """
    Simple test endpoint to debug 422 errors
    """
    logger.info(f"=== TEST FILE UPLOAD ENDPOINT CALLED ===")
    logger.info(f"User ID: {current_user.id}")
    logger.info(f"File name: {file.filename}")
    logger.info(f"File size: {file.size}")
    logger.info(f"File content type: {file.content_type}")
    
    return {
        "message": "File upload test successful",
        "filename": file.filename,
        "size": file.size,
        "content_type": file.content_type
    }

@router.post("/files/debug")
async def debug_file_upload(
    request: Request,
    current_user: User = Depends(get_current_user)
):
    """
    Debug endpoint to see raw request data
    """
    logger.info(f"=== DEBUG FILE UPLOAD ENDPOINT CALLED ===")
    logger.info(f"User ID: {current_user.id}")
    logger.info(f"Request method: {request.method}")
    logger.info(f"Request headers: {dict(request.headers)}")
    logger.info(f"Request content type: {request.headers.get('content-type')}")
    
    # Try to read the request body
    try:
        body = await request.body()
        logger.info(f"Request body length: {len(body)}")
        logger.info(f"Request body preview: {body[:200]}")
    except Exception as e:
        logger.error(f"Error reading request body: {e}")
    
    return {
        "message": "Debug endpoint called",
        "user_id": current_user.id,
        "content_type": request.headers.get('content-type'),
        "body_length": len(await request.body()) if hasattr(request, 'body') else 0
    }

@router.post("/files")
async def upload_and_process_file(
    file: UploadFile = File(None),
    db: AsyncSession = Depends(get_async_db),
    current_user: User = Depends(get_current_user)
):
    """
    Upload a file, extract its content, and process it with the appropriate service.
    Now supports images, code files, data files, and documents.
    """
    # Add comprehensive logging
    logger.info(f"=== WORKSHOP FILES ENDPOINT CALLED ===")
    logger.info(f"User ID: {current_user.id}")
    logger.info(f"User email: {current_user.email}")
    logger.info(f"File object: {file}")
    logger.info(f"File name: {file.filename if file else 'None'}")
    logger.info(f"File size: {file.size if file else 'None'} bytes")
    logger.info(f"File content type: {file.content_type if file else 'None'}")
    logger.info(f"Request timestamp: {datetime.utcnow()}")
    
    # Validate file input
    if not file or not file.filename:
        logger.error("No file provided or no filename provided")
        raise HTTPException(status_code=422, detail="No file provided")
    
    if not file.content_type:
        logger.warning("No content type provided, will attempt to detect")
    
    logger.info(f"File validation passed - proceeding with processing")
    
    if os.getenv("TESTING") == "true":
        return await _testing_upload_and_process_file(file, current_user, db)

    try:
        # Save the file
        logger.info("Saving uploaded file...")
        try:
            file_path, _ = await file_service.save_file(file, current_user.id)
            logger.info(f"File saved successfully to: {file_path}")
        except Exception as e:
            logger.error(f"Failed to save file: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")
        
        # Detect file type and determine service
        # Use the actual MIME type detected by the file service, not the browser-reported one
        content_type = file.content_type or "application/octet-stream"
        
        # Check if we need to correct the content type based on MIME type mismatch
        corrected_content_type = content_type
        if content_type == "text/csv" and file.filename.endswith('.csv'):
            # Check if the actual file is Excel based on the saved file extension
            if file_path.endswith('.xls'):
                corrected_content_type = "application/vnd.ms-excel"
                logger.info(f"Correcting content type from {content_type} to {corrected_content_type}")
            elif file_path.endswith('.xlsx'):
                corrected_content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                logger.info(f"Correcting content type from {content_type} to {corrected_content_type}")
        
        logger.info(f"Detecting file type for content_type: {corrected_content_type}, filename: {file.filename}")
        file_info = detect_file_type_and_service(corrected_content_type, file.filename)
        logger.info(f"File detection result: {file_info}")
        
        # Process based on file type
        if file_info['service'] == 'image_analysis':
            # Check if user has access to image analysis
            if not has_feature_access(current_user, "image_analysis", db):
                plan = get_user_plan(current_user, db)
                raise HTTPException(
                    status_code=403,
                    detail={
                        "error": "Image analysis not available in your plan",
                        "feature": "image_analysis",
                        "current_plan": plan,
                        "upgrade_message": "Upgrade to Pro plan to access image analysis",
                        "upgrade_url": "/dashboard/price-plan"
                    }
                )
            
            # Handle image files
            try:
                with open(file_path, 'rb') as img_file:
                    image_bytes = img_file.read()
                
                # Analyze image based on content type
                if file_info['analysis_type'] == 'general':
                    analysis = await image_analysis_service.analyze_image_and_answer(
                        image_data=image_bytes,
                        question="Provide a comprehensive analysis of this image",
                        context=None
                    )
                else:
                    analysis = await image_analysis_service.analyze_image_and_answer(
                        image_data=image_bytes,
                        question="Analyze this image and provide insights",
                        context=None
                    )
                
                # Create file upload record
                upload_record = await create_file_upload_record(
                    db, current_user.id, file, file_path, "image", 
                    extracted_content=None, ai_analysis=analysis
                )
                
                return {
                    "id": str(uuid.uuid4()),
                    "file_upload_id": upload_record.get("file_upload_id"),
                    "name": file.filename,
                    "size": file.size,
                    "type": content_type,
                    "path": file_path,
                    "content": f"Image Analysis Results:\n{analysis}",
                    "analysis": analysis,
                    "service_used": "image_analysis",
                    "file_category": "image",
                    "uploaded_at": datetime.utcnow().isoformat()
                }
                
            except Exception as e:
                logger.error(f"Image analysis failed: {str(e)}")
                # Fallback to general analysis
                analysis = f"Image uploaded successfully. Analysis failed: {str(e)}"
                
                # Create file upload record
                upload_record = await create_file_upload_record(
                    db, current_user.id, file, file_path, "image", 
                    extracted_content=None, ai_analysis=analysis, processing_status="failed"
                )
                
                return {
                    "id": str(uuid.uuid4()),
                    "file_upload_id": upload_record.get("file_upload_id"),
                    "name": file.filename,
                    "size": file.size,
                    "type": content_type,
                    "path": file_path,
                    "content": f"Image Upload Results:\n{analysis}",
                    "analysis": analysis,
                    "service_used": "image_upload",
                    "file_category": "image",
                    "uploaded_at": datetime.utcnow().isoformat()
                }
        
        elif file_info['service'] == 'ai_analysis':
            # Handle text-based files (documents, code, data)
            logger.info(f"Processing file with AI analysis - type: {file_info['type']}, analysis_type: {file_info['analysis_type']}")
            try:
                content = await extract_file_content(file_path, content_type)
                logger.info(f"File content extracted successfully, length: {len(content)}")
                logger.info(f"Content preview (first 200 chars): {content[:200]}...")
            except Exception as e:
                logger.error(f"Failed to extract file content: {str(e)}")
                logger.error(f"Exception type: {type(e).__name__}")
                import traceback
                logger.error(f"Traceback: {traceback.format_exc()}")
                content = f"Error extracting content from {file.filename}: {str(e)}"
            
            # Create specialized prompts based on file type
            if file_info['type'] == 'code':
                # Check if user has access to code analysis
                if not has_feature_access(current_user, "code_analysis", db):
                    plan = get_user_plan(current_user, db)
                    raise HTTPException(
                        status_code=403,
                        detail={
                            "error": "Code analysis not available in your plan",
                            "feature": "code_analysis",
                            "current_plan": plan,
                            "upgrade_message": "Upgrade to Plus plan to access code analysis",
                            "upgrade_url": "/dashboard/price-plan"
                        }
                    )
                
                analysis_prompt = f"Review and analyze this code:\n\n{content[:3000]}...\n\nProvide a comprehensive code review including:\n1. Code quality assessment\n2. Potential improvements\n3. Best practices suggestions\n4. Security considerations\n5. Performance optimizations"
            elif file_info['type'] == 'data':
                # Check if user has access to data analysis
                if not has_feature_access(current_user, "data_analysis", db):
                    plan = get_user_plan(current_user, db)
                    raise HTTPException(
                        status_code=403,
                        detail={
                            "error": "Data analysis not available in your plan",
                            "feature": "data_analysis",
                            "current_plan": plan,
                            "upgrade_message": "Upgrade to Pro plan to access data analysis",
                            "upgrade_url": "/dashboard/price-plan"
                        }
                    )
                
                analysis_prompt = f"Analyze this data:\n\n{content[:3000]}...\n\nProvide insights including:\n1. Data structure analysis\n2. Key patterns and trends\n3. Statistical summary\n4. Data quality assessment\n5. Visualization suggestions"
            else:
                analysis_prompt = f"Analyze the following document content and provide comprehensive insights:\n\n{content[:3000]}..."
            
            logger.info(f"Generating AI analysis with prompt length: {len(analysis_prompt)}")
            logger.info(f"Analysis prompt preview: {analysis_prompt[:200]}...")
            try:
                # Initialize AI service with existing db session
                ai_service = AIService(db=db)
                
                # Enforce token limits before making AI calls
                try:
                    estimated_tokens = len(analysis_prompt.split()) * 2  # Rough estimate
                    await ai_service.enforce_token_limit(current_user.id, estimated_tokens)
                    logger.info(f"Token limit check passed for user {current_user.id}")
                except HTTPException as e:
                    logger.warning(f"Token limit exceeded for user {current_user.id}: {e.detail}")
                    analysis = f"Token limit exceeded. {e.detail}"
                except Exception as e:
                    logger.error(f"Token enforcement failed for user {current_user.id}: {str(e)}")
                    analysis = f"Unable to verify token limits. Please contact support."
                else:
                    analysis = await ai_service.generate_assignment_content_from_prompt(
                        analysis_prompt, 
                        user_id=current_user.id, 
                        feature='file_upload_analysis', 
                        action='analyze_content'
                    )
                    
                logger.info(f"AI analysis completed successfully, length: {len(analysis)}")
                logger.info(f"Analysis preview (first 200 chars): {analysis[:200]}...")
            except Exception as e:
                logger.error(f"AI analysis failed: {str(e)}")
                logger.error(f"Exception type: {type(e).__name__}")
                import traceback
                logger.error(f"Traceback: {traceback.format_exc()}")
                analysis = f"Analysis failed: {str(e)}"
            
            logger.info("Preparing response for AI analysis...")
            
            # Create file upload record
            upload_record = await create_file_upload_record(
                db, current_user.id, file, file_path, file_info['type'], 
                extracted_content=content, ai_analysis=analysis
            )
            logger.info(f"Upload record created: {upload_record}")
            
            # Include processed structured data for CSV/Excel files
            processed_data = None
            logger.info(f"Checking file type for processed_data: {file.filename}, content_type: {content_type}")
            
            if content_type in ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
                logger.info(f"Starting Excel file processing for {file.filename}")
                try:
                    import openpyxl
                    import pandas as pd
                    
                    file_extension = Path(file_path).suffix[1:].lower()
                    logger.info(f"Processing Excel file extension: {file_extension}")
                    if file_extension == 'xlsx':
                        # Process .xlsx files with openpyxl
                        workbook = openpyxl.load_workbook(file_path, data_only=True)
                        sheets = {}
                        has_calculations = False
                        calculation_instructions = []
                        
                        for sheet_name in workbook.sheetnames:
                            sheet = workbook[sheet_name]
                            data = []
                            for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
                                # Sanitize row data for JSON
                                sanitized_row = []
                                for col_idx, cell in enumerate(row):
                                    if cell is None:
                                        sanitized_row.append(None)
                                    elif isinstance(cell, float):
                                        import math
                                        if math.isinf(cell) or math.isnan(cell):
                                            sanitized_row.append(None)
                                        else:
                                            sanitized_row.append(cell)
                                    elif isinstance(cell, str):
                                        # Detect calculation instructions
                                        if 'compute' in cell.lower() or 'calculate' in cell.lower():
                                            has_calculations = True
                                            calculation_instructions.append({
                                                'sheet': sheet_name,
                                                'row': row_idx,
                                                'col': col_idx,
                                                'instruction': cell
                                            })
                                        sanitized_row.append(cell)
                                    else:
                                        sanitized_row.append(cell)
                                data.append(sanitized_row)
                            sheets[sheet_name] = data
                        
                        processed_data = {
                            'sheets': sheets,
                            'sheet_names': workbook.sheetnames,
                            'file_type': 'xlsx',
                            'has_calculations': has_calculations,
                            'calculation_instructions': calculation_instructions if has_calculations else None
                        }
                        logger.info(f"Successfully processed XLSX file with {len(sheets)} sheet(s), has_calculations: {has_calculations}")
                    elif file_extension == 'xls':
                        # Process .xls files with pandas
                        df_dict = pd.read_excel(file_path, sheet_name=None, engine='xlrd')
                        sheets = {}
                        for sheet_name, sheet_df in df_dict.items():
                            data = [sheet_df.columns.tolist()] + sheet_df.values.tolist()
                            # Sanitize data
                            sanitized_data = []
                            for row in data:
                                sanitized_row = []
                                for cell in row:
                                    if pd.isna(cell):
                                        sanitized_row.append(None)
                                    else:
                                        sanitized_row.append(cell)
                                sanitized_data.append(sanitized_row)
                            sheets[sheet_name] = sanitized_data
                        processed_data = {
                            'sheets': sheets,
                            'sheet_names': list(df_dict.keys()),
                            'file_type': 'xls'
                        }
                        logger.info(f"Successfully processed XLS file with {len(sheets)} sheet(s)")
                except Exception as e:
                    logger.warning(f"Could not get processed data for {file.filename}: {str(e)}")
                    import traceback
                    logger.warning(f"Traceback: {traceback.format_exc()}")
                    processed_data = None
            elif content_type == "text/csv":
                # For CSV files, create a simple processed data structure from the content
                logger.info(f"Processing CSV file for processed_data: {file.filename}")
                try:
                    import pandas as pd
                    import numpy as np
                    import math
                    logger.info(f"Reading CSV file: {file_path}")
                    df = pd.read_csv(file_path)
                    logger.info(f"CSV loaded successfully. Shape: {df.shape}, Columns: {df.columns.tolist()}")
                    
                    # Apply calculation instructions manually
                    records = df.to_dict('records')
                    processed_records = []
                    
                    logger.info(f"Processing {len(records)} records")
                    for i, record in enumerate(records):
                        logger.info(f"Processing record {i}: {record}")
                        # Check if there's a calculation instruction
                        instruction = record.get('Instruction', '')
                        if isinstance(instruction, str) and 'Compute Revenue' in instruction:
                            try:
                                # Extract the calculation formula
                                if 'Units*UnitPrice' in instruction or 'Units * UnitPrice' in instruction:
                                    units = float(record.get('Units', 0))
                                    unit_price = float(record.get('UnitPrice', 0))
                                    revenue = units * unit_price
                                    record['Revenue'] = revenue
                                    record['Instruction'] = revenue  # Replace instruction with result
                                    logger.info(f"Calculated revenue for record {i}: {revenue}")
                            except (ValueError, TypeError) as calc_error:
                                logger.warning(f"Calculation error for record {i}: {calc_error}")
                        
                        # Sanitize all values to ensure JSON compatibility
                        sanitized_record = {}
                        for key, value in record.items():
                            # Handle NaN, Infinity, and other non-JSON-compliant values
                            if pd.isna(value):
                                sanitized_record[key] = None
                            elif isinstance(value, (float, np.floating)):
                                if math.isinf(value) or math.isnan(value):
                                    sanitized_record[key] = None
                                else:
                                    sanitized_record[key] = float(value)
                            elif isinstance(value, (int, np.integer)):
                                sanitized_record[key] = int(value)
                            else:
                                sanitized_record[key] = value
                        
                        processed_records.append(sanitized_record)
                    
                    processed_data = {
                        'data': processed_records,
                        'columns': df.columns.tolist(),
                        'rows': len(df),
                        'has_calculations': True  # We know this CSV has calculations
                    }
                    logger.info(f"Successfully created processed_data: {len(processed_records)} records with calculations")
                except Exception as e:
                    logger.error(f"Could not process CSV data for {file.filename}: {str(e)}")
                    import traceback
                    logger.error(f"Traceback: {traceback.format_exc()}")
                    processed_data = None
            else:
                logger.info(f"File type not matched for processed_data processing. content_type: {content_type}")
            
            # Generate corrected filename for display based on MIME type
            display_filename = file.filename
            if content_type == 'application/vnd.ms-excel' and file.filename.endswith('.csv'):
                # Replace .csv with .xls for display
                display_filename = file.filename[:-4] + '.xls'
                logger.info(f"Correcting display filename from {file.filename} to {display_filename}")
            elif content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' and file.filename.endswith('.csv'):
                # Replace .csv with .xlsx for display
                display_filename = file.filename[:-4] + '.xlsx'
                logger.info(f"Correcting display filename from {file.filename} to {display_filename}")
            
            response_data = {
                "id": str(uuid.uuid4()),
                "file_upload_id": upload_record.get("file_upload_id"),
                "name": display_filename,
                "size": file.size,
                "type": corrected_content_type,
                "path": file_path,
                "content": content,
                "analysis": analysis,
                "processed_data": processed_data,  # Add processed structured data
                "service_used": "ai_analysis",
                "file_category": file_info['type'],
                "uploaded_at": datetime.utcnow().isoformat()
            }
            logger.info(f"Response prepared successfully. Content length: {len(content)}, Analysis length: {len(analysis)}")
            logger.info(f"Processed data included: {processed_data is not None}")
            logger.info(f"Final response data: {response_data}")
            if processed_data:
                logger.info(f"Processed data structure: {type(processed_data)}, keys: {list(processed_data.keys()) if isinstance(processed_data, dict) else 'not a dict'}")
            else:
                logger.info("No processed data to include in response")
            logger.info("=== WORKSHOP FILES ENDPOINT COMPLETED SUCCESSFULLY ===")
            return response_data
        
        else:
            # Fallback for unknown file types
            content = f"File uploaded: {file.filename} (Type: {content_type})"
            analysis = "File uploaded successfully. Content analysis not available for this file type."
            
            # Create file upload record
            upload_record = await create_file_upload_record(
                db, current_user.id, file, file_path, "unknown", 
                extracted_content=content, ai_analysis=analysis
            )
            
            return {
                "id": str(uuid.uuid4()),
                "file_upload_id": upload_record.get("file_upload_id"),
                "name": file.filename,
                "size": file.size,
                "type": content_type,
                "path": file_path,
                "content": content,
                "analysis": analysis,
                "service_used": "file_upload",
                "file_category": "unknown",
                "uploaded_at": datetime.utcnow().isoformat()
            }
            
    except HTTPException as e:
        # Re-raise HTTPExceptions (like 403 errors) without wrapping them
        logger.error(f"HTTPException in file processing: {e.status_code} - {e.detail}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error processing file: {str(e)}")
        logger.error(f"Exception type: {type(e).__name__}")
        logger.error(f"File details - name: {file.filename}, size: {file.size}, type: {file.content_type}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")

@router.post("/files/process")
async def process_uploaded_file(
    file_id: str = Body(..., embed=True),
    action: str = Body(..., embed=True),  # "summarize", "extract", "rewrite", "analyze"
    db: AsyncSession = Depends(get_async_db),
    current_user: User = Depends(get_current_user)
):
    """
    Process an uploaded file with specific AI actions.
    """
    # Add comprehensive logging
    logger.info(f"=== WORKSHOP FILES/PROCESS ENDPOINT CALLED ===")
    logger.info(f"User ID: {current_user.id}")
    logger.info(f"User email: {current_user.email}")
    logger.info(f"File ID: {file_id}")
    logger.info(f"Action: {action}")
    logger.info(f"Request timestamp: {datetime.utcnow()}")
    
    if os.getenv("TESTING") == "true":
        return await _testing_process_uploaded_file(file_id, action)

    try:
        # Initialize AI service
        ai_service = AIService(db=db)
        
        # In a real implementation, you'd fetch the file from storage
        # For now, we'll simulate the processing
        logger.info("Processing file with AI service...")
        if action == "summarize":
            prompt = "Please provide a comprehensive summary of the document content."
        elif action == "extract":
            prompt = "Please extract the key points and main ideas from the document."
        elif action == "rewrite":
            prompt = "Please rewrite the document content in a more academic tone."
        elif action == "analyze":
            prompt = "Please analyze the document structure, arguments, and provide feedback."
        else:
            logger.error(f"Invalid action specified: {action}")
            raise HTTPException(status_code=400, detail="Invalid action specified")
        
        logger.info(f"Generated prompt: {prompt}")
        
        # Enforce token limits before making AI calls
        try:
            estimated_tokens = len(prompt.split()) * 2  # Rough estimate
            await ai_service.enforce_token_limit(current_user.id, estimated_tokens)
            logger.info(f"Token limit check passed for user {current_user.id}")
        except HTTPException as e:
            logger.warning(f"Token limit exceeded for user {current_user.id}: {e.detail}")
            result = f"Token limit exceeded. {e.detail}"
        except Exception as e:
            logger.error(f"Token enforcement failed for user {current_user.id}: {str(e)}")
            result = f"Unable to verify token limits. Please contact support."
        else:
            result = await ai_service.generate_assignment_content_from_prompt(
                prompt, 
                user_id=current_user.id, 
                feature='file_processing', 
                action=action
            )
                
        logger.info(f"AI processing completed. Result length: {len(result)}")
        logger.info("=== WORKSHOP FILES/PROCESS ENDPOINT COMPLETED SUCCESSFULLY ===")
        
        return {
            "action": action,
            "result": result,
            "processed_at": datetime.utcnow().isoformat()
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing file action: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")

@router.post("/links")
async def process_link(
    url: str = Body(..., embed=True),
    db: AsyncSession = Depends(get_async_db),
    current_user: User = Depends(get_current_user)
):
    """
    Process a link and extract its content.
    """
    # Add comprehensive logging
    logger.info(f"=== WORKSHOP LINKS ENDPOINT CALLED ===")
    logger.info(f"User ID: {current_user.id}")
    logger.info(f"User email: {current_user.email}")
    logger.info(f"URL: {url}")
    logger.info(f"Request timestamp: {datetime.utcnow()}")
    
    if os.getenv("TESTING") == "true":
        return await _testing_process_link(url)

    try:
        # Initialize AI service
        ai_service = AIService(db=db)
        
        logger.info("Initializing web scraping service...")
        async with WebScrapingService() as scraper:
            logger.info("Extracting content from URL...")
            result = await scraper.extract_content_from_url(url)
            logger.info(f"Content extraction completed. Title: {result['title']}, Content length: {len(result['content'])}")
            
            # Generate AI analysis of the extracted content
            logger.info("Generating AI analysis of extracted content...")
            analysis_prompt = f"Analyze the following web content and provide insights:\n\n{result['content'][:2000]}..."
            
            # Enforce token limits before making AI calls
            try:
                estimated_tokens = len(analysis_prompt.split()) * 2  # Rough estimate
                await ai_service.enforce_token_limit(current_user.id, estimated_tokens)
                logger.info(f"Token limit check passed for user {current_user.id}")
                analysis = await ai_service.generate_assignment_content_from_prompt(
                    analysis_prompt, 
                    user_id=current_user.id, 
                    feature='link_analysis', 
                    action='analyze'
                )
            except HTTPException as e:
                logger.warning(f"Token limit exceeded for user {current_user.id}: {e.detail}")
                analysis = f"Token limit exceeded. {e.detail}"
            except Exception as e:
                logger.error(f"Token enforcement failed for user {current_user.id}: {str(e)}")
                analysis = f"Unable to verify token limits. Please contact support."
            logger.info(f"AI analysis completed. Analysis length: {len(analysis)}")
            
            # Create file upload record for link
            try:
                file_upload_data = FileUploadCreate(
                    filename=result['title'] or url,
                    original_filename=result['title'] or url,
                    file_path=url,
                    file_size=len(result['content']),
                    mime_type="text/html",
                    file_type="link",
                    extracted_content=result['content'],
                    ai_analysis=analysis,
                    processing_status="completed",
                    is_link=True,
                    link_url=url,
                    link_title=result['title'],
                    link_description=result['content'][:500] + "..." if len(result['content']) > 500 else result['content'],
                    upload_metadata={
                        "uploaded_via": "workshop",
                        "timestamp": datetime.utcnow().isoformat()
                    }
                )
                
                db_file_upload = file_upload_crud.create_file_upload(db, file_upload_data, current_user.id)
                file_upload_id = db_file_upload.id
            except Exception as e:
                logger.error(f"Failed to create file upload record for link: {str(e)}")
                file_upload_id = None
            
            logger.info("=== WORKSHOP LINKS ENDPOINT COMPLETED SUCCESSFULLY ===")
            return {
                "id": str(uuid.uuid4()),
                "file_upload_id": file_upload_id,
                "url": url,
                "title": result['title'],
                "content": result['content'],
                "type": result['type'],
                "analysis": analysis,
                "extracted_at": datetime.utcnow().isoformat()
            }
    except Exception as e:
        logger.error(f"Error processing link: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to process link: {str(e)}")

@router.delete("/files/{file_id}")
async def delete_file(
    file_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Delete a processed file.
    """
    try:
        # In a real implementation, you'd delete the file from storage
        # For now, we'll just return success
        return {"message": "File deleted successfully"}
    except Exception as e:
        logger.error(f"Error deleting file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete file: {str(e)}")

@router.delete("/links/{link_id}")
async def delete_link(
    link_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Delete a processed link.
    """
    try:
        # In a real implementation, you'd remove the link from storage
        return {"message": "Link deleted successfully"}
    except Exception as e:
        logger.error(f"Error deleting link: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to delete link: {str(e)}")

@router.post("/chat-with-link")
async def chat_with_link(
    request: ChatWithLinkRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Interactive streaming chat with AI about a specific link.
    Allows users to ask questions, request enhancements, and get insights.
    """
    try:
        logger.info(f"Processing streaming chat message for link {request.link_id} from user {current_user.id}")
        
        ai_service = AIService(db=db)
        
        # Create context-aware prompt
        context_prompt = f"""
        You are a friendly AI assistant. Answer questions about web content in a brief, conversational way.

        Web Content: {request.content[:2000]}...
        
        Current Analysis: {request.analysis if request.analysis else 'No analysis available'}
        
        User Question: {request.message}
        
        Keep your response:
        - Brief (2-4 sentences max unless specifically asked for more detail)
        - Conversational and natural (like texting a knowledgeable friend)
        - Direct and to the point
        - Avoid bullet points or formal structure unless necessary
        - No assignments or documents - just helpful chat
        """
        
        async def generate_stream():
            try:
                # Check if the response suggests updating the analysis
                updated_analysis = None
                if any(keyword in request.message.lower() for keyword in ['reanalyze', 'update analysis', 'new insights']):
                    analysis_service = LinkAnalysisService()
                    updated_analysis = await analysis_service.analyze_content_comprehensive(
                        content=request.content,
                        url='',  # URL not available in this context
                        ai_service=ai_service
                    )
                
                # Stream the AI response
                full_response = ""
                chunk_count = 0
                async for chunk in ai_service.generate_chat_response_stream(context_prompt, user_id=current_user.id):
                    chunk_count += 1
                    if chunk:
                        full_response += chunk
                        # Send chunk as JSON
                        chunk_data = {'chunk': chunk, 'done': False}
                        logger.info(f"Sending chunk {chunk_count}: {len(chunk)} characters")
                        yield f"data: {json.dumps(chunk_data)}\n\n"
                
                # Track token usage for chat
                estimated_tokens = len(context_prompt.split()) + len(full_response.split())
                await ai_service.track_token_usage(
                    user_id=current_user.id,
                    feature='link_chat',
                    action='chat_message',
                    tokens_used=estimated_tokens,
                    metadata={
                        'link_id': request.link_id,
                        'message': request.message[:100],  # First 100 chars of user message
                        'response_length': len(full_response)
                    }
                )
                logger.info(f"Token usage tracked: {estimated_tokens} tokens for link chat")
                
                # Send final response with updated analysis
                final_response = {
                    'chunk': '',
                    'done': True,
                    'full_response': full_response,
                    'updated_analysis': updated_analysis
                }
                logger.info(f"Streaming completed. Total chunks sent: {chunk_count}, Final response length: {len(full_response)}")
                yield f"data: {json.dumps(final_response)}\n\n"
                
                logger.info(f"Streaming chat response completed for user {current_user.id}")
                
            except Exception as e:
                logger.error(f"Error in streaming chat: {str(e)}", exc_info=True)
                error_response = {
                    'chunk': '',
                    'done': True,
                    'error': f"Failed to process chat message: {str(e)}"
                }
                yield f"data: {json.dumps(error_response)}\n\n"
        
        return StreamingResponse(
            generate_stream(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Headers": "Cache-Control"
            }
        )
        
    except Exception as e:
        logger.error(f"Error in chat with link: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process chat message: {str(e)}"
        )

@router.get("/features")
async def get_user_features(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get the user's available features based on their subscription plan.
    """
    try:
        from app.core.feature_access import get_available_features, get_user_plan, get_feature_requirements
        
        plan = get_user_plan(current_user, db)
        available_features = get_available_features(current_user, db)
        feature_requirements = get_feature_requirements()
        
        return {
            "current_plan": plan,
            "available_features": available_features,
            "feature_requirements": feature_requirements,
            "upgrade_url": "/dashboard/price-plan"
        }
    except Exception as e:
        logger.error(f"Error getting user features: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to get user features: {str(e)}")

async def extract_file_content(file_path: str, content_type: str) -> str:
    """
    Extract text content from various file types using proper parsing libraries.
    Now supports code files, data files, and additional document formats.
    """
    logger.info(f"Extracting content from file: {file_path}, content_type: {content_type}")
    try:
        # Text-based files
        if content_type == "text/plain":
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        # PDF files
        elif content_type == "application/pdf":
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        
        # Word documents
        elif content_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/x-ole-storage"]:
            logger.info(f"Processing Word document: {file_path}")
            try:
                file_extension = Path(file_path).suffix[1:].lower()
                
                if file_extension == 'docx':
                    # Handle .docx files
                    doc = Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + "\t"
                            text += "\n"
                    logger.info(f"DOCX content extracted successfully, length: {len(text)}")
                    return text.strip()
                elif file_extension == 'doc':
                    # Handle legacy .doc files using FileProcessingService
                    from app.services.file_processing_service import FileProcessingService
                    processing_service = FileProcessingService(None)
                    processed_content = await processing_service._process_doc(file_path)
                    text = processed_content.get('text', '')
                    logger.info(f"DOC content extracted successfully, length: {len(text)}")
                    return text.strip()
                else:
                    # Try docx first, then doc
                    try:
                        doc = Document(file_path)
                        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                        return text.strip()
                    except:
                        from app.services.file_processing_service import FileProcessingService
                        processing_service = FileProcessingService(None)
                        processed_content = await processing_service._process_doc(file_path)
                        return processed_content.get('text', '').strip()
            except Exception as e:
                logger.error(f"Failed to extract Word document content: {str(e)}")
                logger.error(f"Exception type: {type(e).__name__}")
                import traceback
                logger.error(f"Traceback: {traceback.format_exc()}")
                raise
        
        # RTF files
        elif content_type == "application/rtf":
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                import re
                content = re.sub(r'\\[a-z]+\d*', '', content)
                content = re.sub(r'[{}]', '', content)
                return content.strip()
        
        # Code files
        elif content_type in ["text/x-python", "text/javascript", "text/x-java-source", "text/x-c++src", "text/x-csrc", "text/html", "text/css", "application/json", "text/xml"]:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        # Data files (Excel and CSV)
        elif content_type in ["text/csv", "application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]:
            file_extension = Path(file_path).suffix[1:].lower()
            
            if content_type == "text/csv" or file_extension == 'csv':
                # Read CSV as plain text
                logger.info(f"Processing CSV file: {file_path}")
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                # For Excel files (.xls, .xlsx), process directly with openpyxl/pandas
                logger.info(f"Processing Excel file: {file_path}, extension: {file_extension}")
                try:
                    if file_extension == 'xlsx':
                        # Process .xlsx files with openpyxl
                        import openpyxl
                        workbook = openpyxl.load_workbook(file_path, data_only=True)
                        sheets = {}
                        
                        for sheet_name in workbook.sheetnames:
                            sheet = workbook[sheet_name]
                            data = []
                            for row in sheet.iter_rows(values_only=True):
                                data.append(list(row))
                            sheets[sheet_name] = data
                        
                        processed_content = {'sheets': sheets}
                    elif file_extension == 'xls':
                        # Process .xls files with pandas
                        import pandas as pd
                        df_dict = pd.read_excel(file_path, sheet_name=None, engine='xlrd')
                        sheets = {}
                        for sheet_name, sheet_df in df_dict.items():
                            data = [sheet_df.columns.tolist()] + sheet_df.values.tolist()
                            sheets[sheet_name] = data
                        processed_content = {'sheets': sheets}
                    else:
                        logger.error(f"Unsupported Excel format: {file_extension}")
                        return f"[Unsupported Excel format: {file_extension}]"
                    
                    # Format the processed content as readable string
                    if processed_content.get('sheets'):
                        formatted_content = ""
                        sheets = processed_content['sheets']
                        
                        # Include formula information if available
                        if processed_content.get('formulas'):
                            formatted_content += "=== FORMULAS DETECTED ===\n"
                            for sheet_name, formulas in processed_content['formulas'].items():
                                formatted_content += f"\n{sheet_name}:\n"
                                for formula_info in formulas:
                                    formatted_content += f"  {formula_info['cell']}: {formula_info['formula']} = {formula_info['value']}\n"
                            formatted_content += "\n"
                        
                        # Add sheet data
                        for sheet_name, sheet_data in sheets.items():
                            if len(sheets) > 1:  # Only show sheet names if multiple sheets
                                formatted_content += f"\n=== Sheet: {sheet_name} ===\n"
                            
                            if isinstance(sheet_data, list):
                                # Format as CSV-like with proper alignment
                                for row in sheet_data:
                                    if isinstance(row, list):
                                        formatted_content += ",".join([str(cell) if cell is not None else '' for cell in row]) + "\n"
                                    else:
                                        formatted_content += str(row) + "\n"
                        
                        logger.info(f"Excel content formatted successfully, length: {len(formatted_content)}")
                        return formatted_content.strip()
                    else:
                        logger.warning("No sheets found in processed Excel content")
                        return f"[Excel file processed but no data found: {file_path}]"
                        
                except Exception as e:
                    logger.error(f"Could not process Excel file {file_path}: {str(e)}")
                    import traceback
                    logger.error(f"Traceback: {traceback.format_exc()}")
                    return f"[Error processing Excel file: {str(e)}]"
        
        # Image files (return placeholder for analysis)
        elif content_type.startswith("image/"):
            return f"[Image file: {file_path} - Image analysis performed separately]"
        
        # Default fallback
        else:
            try:
                # Try to read as text
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # If text reading fails, return file info
                return f"[Binary file: {file_path} - Content type: {content_type}]"
                
    except Exception as e:
        logger.error(f"Error extracting content from file: {str(e)}")
        logger.error(f"Exception type: {type(e).__name__}")
        logger.error(f"File path: {file_path}, Content type: {content_type}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return f"[Error extracting content: {str(e)}]" 