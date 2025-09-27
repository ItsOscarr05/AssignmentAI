#!/usr/bin/env python3
"""
API Endpoint Testing Script
This script tests the file processing API endpoints directly.
"""
import requests
import json
import os
from pathlib import Path

class FileCompletionAPITester:
    def __init__(self, base_url="http://localhost:8000"):
        self.base_url = base_url
        self.token = None
        
    def login(self, email, password):
        """Login and get authentication token"""
        url = f"{self.base_url}/api/v1/auth/login"
        data = {
            "email": email,
            "password": password
        }
        
        try:
            response = requests.post(url, json=data)
            if response.status_code == 200:
                result = response.json()
                self.token = result.get("access_token")
                print(f"✅ Login successful!")
                return True
            else:
                print(f"❌ Login failed: {response.status_code} - {response.text}")
                return False
        except Exception as e:
            print(f"❌ Login error: {e}")
            return False
    
    def test_file_upload_and_fill(self, file_path):
        """Test file upload and filling"""
        if not self.token:
            print("❌ Not logged in. Please login first.")
            return False
        
        url = f"{self.base_url}/api/v1/file-processing/fill"
        headers = {
            "Authorization": f"Bearer {self.token}"
        }
        
        try:
            with open(file_path, 'rb') as file:
                files = {'file': file}
                response = requests.post(url, headers=headers, files=files)
            
            if response.status_code == 200:
                result = response.json()
                print(f"✅ File processing successful!")
                print(f"📁 Original file: {result.get('file_name')}")
                print(f"📁 Completed file: {result.get('filled_file_name')}")
                print(f"📊 Sections filled: {result.get('sections_filled', 0)}")
                print(f"📄 File type: {result.get('file_type')}")
                return result
            else:
                print(f"❌ File processing failed: {response.status_code}")
                print(f"Error: {response.text}")
                return None
                
        except Exception as e:
            print(f"❌ File processing error: {e}")
            return None
    
    def test_file_analysis(self, file_path):
        """Test file analysis only"""
        if not self.token:
            print("❌ Not logged in. Please login first.")
            return False
        
        url = f"{self.base_url}/api/v1/file-processing/analyze"
        headers = {
            "Authorization": f"Bearer {self.token}"
        }
        
        try:
            with open(file_path, 'rb') as file:
                files = {'file': file}
                response = requests.post(url, headers=headers, files=files)
            
            if response.status_code == 200:
                result = response.json()
                print(f"✅ File analysis successful!")
                print(f"📁 File: {result.get('file_name')}")
                print(f"📄 File type: {result.get('file_type')}")
                print(f"🔍 Fillable sections: {len(result.get('fillable_sections', []))}")
                print(f"📝 Analysis: {result.get('analysis', '')[:200]}...")
                return result
            else:
                print(f"❌ File analysis failed: {response.status_code}")
                print(f"Error: {response.text}")
                return None
                
        except Exception as e:
            print(f"❌ File analysis error: {e}")
            return None
    
    def test_supported_formats(self):
        """Test supported formats endpoint"""
        url = f"{self.base_url}/api/v1/file-processing/supported-formats"
        
        try:
            response = requests.get(url)
            if response.status_code == 200:
                result = response.json()
                formats = result.get('supported_formats', [])
                print(f"✅ Supported formats retrieved: {len(formats)} formats")
                
                # Show some examples
                for fmt in formats[:5]:
                    print(f"  - {fmt.get('extension')}: {fmt.get('description')}")
                
                return result
            else:
                print(f"❌ Failed to get supported formats: {response.status_code}")
                return None
                
        except Exception as e:
            print(f"❌ Error getting supported formats: {e}")
            return None

def create_test_file():
    """Create a test file for testing"""
    from docx import Document
    import tempfile
    
    # Create a temporary DOCX file
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test document for file completion.')
    
    # Add a table with placeholders
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = 'Question'
    table.cell(0, 1).text = 'Answer'
    table.cell(1, 0).text = 'What is the purpose?'
    table.cell(1, 1).text = '_____'
    table.cell(2, 0).text = 'What are the benefits?'
    table.cell(2, 1).text = '[INSERT ANSWER HERE]'
    
    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name

def main():
    """Main testing function"""
    print("🧪 File Completion API Tester")
    print("=" * 40)
    
    # Initialize tester
    tester = FileCompletionAPITester()
    
    # Test supported formats first (no auth needed)
    print("\n1. Testing supported formats...")
    tester.test_supported_formats()
    
    # Get login credentials
    print("\n2. Authentication required for other tests...")
    email = input("Enter your email: ").strip()
    password = input("Enter your password: ").strip()
    
    if not tester.login(email, password):
        print("❌ Cannot proceed without authentication")
        return
    
    # Create test file
    print("\n3. Creating test file...")
    test_file_path = create_test_file()
    print(f"✅ Test file created: {test_file_path}")
    
    try:
        # Test file analysis
        print("\n4. Testing file analysis...")
        analysis_result = tester.test_file_analysis(test_file_path)
        
        # Test file completion
        print("\n5. Testing file completion...")
        completion_result = tester.test_file_upload_and_fill(test_file_path)
        
        if completion_result:
            print("\n🎉 All tests completed successfully!")
            print("\nNext steps:")
            print("1. Check the uploads directory for the completed file")
            print("2. Open the completed file to verify it's not corrupted")
            print("3. Compare with the original to see what was filled")
        
    finally:
        # Clean up test file
        if os.path.exists(test_file_path):
            os.unlink(test_file_path)
            print(f"\n🧹 Cleaned up test file: {test_file_path}")

if __name__ == "__main__":
    main()
